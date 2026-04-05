import streamlit as st
st.set_page_config(page_title="School Scheduler Pro", layout="wide", page_icon="📚")

import pandas as pd
import copy, json, os, random, re, zipfile
import sqlite3

from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches  # noqa: F401  (kept for compatibility)

from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet

# ══════════════════════════════════════════════════════════════════════════════
# AI CLIENT SETUP
# ══════════════════════════════════════════════════════════════════════════════

from openai import OpenAI as _OpenAI

# Set the API key BEFORE creating the client
os.environ["OPENAI_API_KEY"] = ""   # 👈 Replace with your actual HF token

_openai_key = os.getenv("OPENAI_API_KEY", "")

openai_client = _OpenAI(
    api_key=_openai_key,
    base_url="https://router.huggingface.co/v1"
) if _openai_key else None

def ask_ai(prompt):
    if not openai_client:
        return "❌ API not connected"

    try:
        response = openai_client.chat.completions.create(
            model="moonshotai/Kimi-K2-Instruct",
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content

    except Exception as e:
        return f"❌ Error: {str(e)}"

# ══════════════════════════════════════════════════════════════════════════════
# CONSTANTS
# ══════════════════════════════════════════════════════════════════════════════

DAYS = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]

MON_WED_TIMES = {
    "P1": "8:00-8:50",   "P2": "8:50-9:30",   "P3": "9:30-10:10",
    "P4": "10:10-10:50", "Lunch": "10:50-11:20",
    "P5": "11:20-12:00", "P6": "12:00-12:40",
    "P7": "12:40-1:20",  "P8": "1:20-2:00",
}
THURSDAY_TIMES = {
    "P1": "8:50-9:30",   "P2": "9:30-10:05",  "P3": "10:05-10:40",
    "P4": "10:40-11:15", "Lunch": "11:15-11:35",
    "P5": "11:35-12:10", "P6": "12:10-12:45",
    "P7": "12:45-1:15",  "P8": "1:15-2:00",
}
FRIDAY_TIMES = {
    "P1": "8:00-8:40",  "P2": "8:40-9:15",   "P3": "9:15-9:50",
    "P4": "9:50-10:25", "Lunch": "10:25-10:55",
    "P5": "10:55-11:35","P6": "11:35-12:15",
}
ALL_PERIODS   = ["P1","P2","P3","P4","Lunch","P5","P6","P7","P8"]
BRAND_BLUE    = "#1f4e79"

ADMIN_USERNAME  = "admin";  ADMIN_PASSWORD  = "Kinz@420"
HEAD_USERNAME   = "head";   HEAD_PASSWORD   = "9999"

VIEWER_USERNAME = "teacher"   # ✅ yeh add karo
VIEWER_PASSWORD = "1234"

SUBJECT_GROUPS = {
    "ENGLISH": ["ENGLISH"],
    "URDU": ["URDU"],
    "ISLAMIAT_ARABIC": ["ISLAMIAT", "ARABIC"],
    "MATH": ["MATH"],
    "SCIENCE": ["GENERAL SCIENCE", "BIOLOGY", "PHYSICS", "CHEMISTRY"],
    "COMPUTER": ["COMPUTER"],
    "SOCIAL": ["HISTORY", "GEOGRAPHY", "SST"],
    "SINDHI": ["SINDHI"],
    "TS": ["TS", "PH", "GP"]
}

# ══════════════════════════════════════════════════════════════════════════════
# DATABASE
# ══════════════════════════════════════════════════════════════════════════════

conn = sqlite3.connect("school.db", check_same_thread=False)
cur  = conn.cursor()
cur.execute("CREATE TABLE IF NOT EXISTS app_data(id INTEGER PRIMARY KEY, data TEXT)")
conn.commit()

# ══════════════════════════════════════════════════════════════════════════════
# CONSTRAINT ENGINE GLOBAL STATE
# ══════════════════════════════════════════════════════════════════════════════

teacher_day_load  = {}
teacher_timeline  = {}
double_used       = {}
subject_remaining = {}
teacher_three_streak_count = {}

# ══════════════════════════════════════════════════════════════════════════════
# CORE HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def clean(x):
    return str(x).strip().upper()
def get_group(subject):
    subject = str(subject).upper()
    for group, items in SUBJECT_GROUPS.items():
        for i in items:
            if i in subject:
                return group
    return "OTHER"
def get_periods(day):
    if day == "Friday":
        return ["P1","P2","P3","P4","Lunch","P5","P6"]
    return ["P1","P2","P3","P4","Lunch","P5","P6","P7","P8"]
def get_time(day, period):
    if day in ("Monday","Tuesday","Wednesday"): return MON_WED_TIMES.get(period,"")
    if day == "Thursday":                       return THURSDAY_TIMES.get(period,"")
    if day == "Friday":                         return FRIDAY_TIMES.get(period,"")
    return ""

def slot(section, day, period):
    return st.session_state.timetable[section][day][period]


def build_subject_timetable(group_name):
    df = build_principal_matrix().copy()

    cols = list(df.columns)
    fixed_cols = ["Day", "P. No.", "Bell Timing"]

    selected_teachers = []

    for teacher in cols:
        if teacher in fixed_cols:
            continue

        subjects = []
        for sec in st.session_state.teacher_assignment.get(teacher, {}):
            subjects += st.session_state.teacher_assignment[teacher][sec]

        if any(get_group(s) == group_name for s in subjects):
            selected_teachers.append(teacher)

    return df[fixed_cols + selected_teachers]

# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE  (single initialisation block)
# ══════════════════════════════════════════════════════════════════════════════

_DEFAULTS = {
    "teachers":{}, "sections":{}, "class_teachers":{},
    "subject_config":{}, "teacher_assignment":{}, "timetable":{},
    "logged_in":False, "role":None,
    "class_view_files": [],
    "teacher_view_files": [],
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

# ══════════════════════════════════════════════════════════════════════════════
# PERSISTENCE
# ══════════════════════════════════════════════════════════════════════════════

def save_all_data():
    data = json.dumps({k: st.session_state[k] for k in
        ("teachers","sections","class_teachers","subject_config",
         "teacher_assignment","timetable")})
    cur.execute("SELECT COUNT(*) FROM app_data")
    if cur.fetchone()[0] == 0:
        cur.execute("INSERT INTO app_data(id,data) VALUES(1,?)", (data,))
    else:
        cur.execute("UPDATE app_data SET data=? WHERE id=1", (data,))
    conn.commit()

def load_all_data():
    cur.execute("SELECT data FROM app_data LIMIT 1")
    row = cur.fetchone()
    if row:
        data = json.loads(row[0])
        for key in ("teachers","sections","class_teachers",
                    "subject_config","teacher_assignment","timetable"):
            st.session_state[key] = data.get(key, {})
    st.session_state.teachers = {clean(t):{} for t in st.session_state.teachers}
    for sec in st.session_state.timetable:
        for day in DAYS:
            for p in get_periods(day):
                cell = st.session_state.timetable[sec][day][p]
                cell["teacher"] = clean(cell.get("teacher",""))

if "data_loaded" not in st.session_state:
    load_all_data()
    st.session_state.data_loaded = True

# ══════════════════════════════════════════════════════════════════════════════
# LOGIN
# ══════════════════════════════════════════════════════════════════════════════

if not st.session_state.logged_in:
    st.markdown("""
    <style>
    #MainMenu, footer, header {visibility: hidden;}
    [data-testid="stAppViewContainer"] {
        background: linear-gradient(135deg, #0d2b4e 0%, #1f4e79 50%, #2e6da4 100%);
        min-height: 100vh;
    }
    [data-testid="stMain"] {background: transparent !important;}
    .login-banner { text-align: center; padding: 48px 20px 20px; }
    .login-banner h1 {
        color: #ffffff; font-size: 2.4rem; font-weight: 800;
        letter-spacing: 0.5px; margin: 0 0 6px;
        text-shadow: 0 2px 8px rgba(0,0,0,.35);
    }
    .login-banner p { color: rgba(255,255,255,0.78); font-size: 1.05rem; margin: 0; }
    .login-card {
        background: #ffffff; border-radius: 18px; padding: 40px 44px 36px;
        max-width: 420px; margin: 22px auto 0;
        box-shadow: 0 24px 60px rgba(0,0,0,0.30), 0 6px 18px rgba(0,0,0,0.18);
    }
    .login-card h3 {
        color: #1f4e79; font-size: 1.25rem; font-weight: 700;
        margin: 0 0 22px; text-align: center;
    }
    .login-footer { text-align: center; color: #8fa8c8; font-size: 0.78rem; margin-top: 18px; }

    @media (max-width: 768px) {
        .login-banner { padding: 28px 12px 14px; }
        .login-banner h1 { font-size: 1.5rem !important; }
        .login-banner p  { font-size: 0.9rem; }
        .login-card {
            padding: 28px 20px 24px;
            margin: 14px 10px 0;
            border-radius: 14px;
        }
    }
    </style>
    <div class="login-banner">
        <h1>🏫 DHACSS Phase IV Campus</h1>
        <p>Academic Timetable Intelligence System</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="login-card"><h3>🔐 Sign In</h3>', unsafe_allow_html=True)
    _, col, _ = st.columns([1,2,1])
    with col:
        username  = st.text_input("Username")
        password  = st.text_input("Password", type="password")
        login_btn = st.button("Login", use_container_width=True)
    st.markdown('<div class="login-footer">DHACSS Phase IV &nbsp;·&nbsp; Timetable Pro v2</div></div>',
                unsafe_allow_html=True)

    if login_btn:
        if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
            role = "admin"

        elif username == HEAD_USERNAME and password == HEAD_PASSWORD:
            role = "viewer"  # ya "head" agar separate role chahte ho

        elif username == VIEWER_USERNAME and password == VIEWER_PASSWORD:
            role = "viewer"

        else:
            role = None

        if role:
            st.session_state.logged_in = True
            st.session_state.role = role
            st.rerun()
        else:
            st.error("Invalid credentials.")

    st.stop()

# ══════════════════════════════════════════════════════════════════════════════
# GLOBAL STYLES
# ══════════════════════════════════════════════════════════════════════════════

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700;800&family=DM+Sans:ital,opsz,wght@0,9..40,300;0,9..40,400;0,9..40,500;1,9..40,400&display=swap');
:root {
    --navy:#0d2f52; --blue-800:#1f4e79; --blue-600:#2b6cb0; --blue-400:#4a90e2;
    --blue-100:#dbeafe; --accent:#f59e0b; --green:#10b981; --red:#ef4444;
    --bg:#f0f4f9; --surface:#ffffff; --border:#e2e8f0; --text:#1e293b;
    --white:#64748b; --radius:13px;
    --shadow:0 4px 24px rgba(31,78,121,0.10);
    --shadow-lg:0 8px 40px rgba(31,78,121,0.18);
}
.login-card div[data-testid="stTextInput"] label,
.login-card div[data-testid="stTextInput"] p {
    color: white !important;
}
*, *::before, *::after { box-sizing: border-box; }
html, body, [class*="css"] { font-family:'DM Sans',sans-serif !important; color:var(--text) !important; }
[data-testid="stAppViewContainer"] { background:var(--bg) !important; }
[data-testid="stMain"]             { background:var(--bg) !important; }
[data-testid="stMain"] > div { padding-left: 1rem !important; padding-right: 1rem !important; }
[data-testid="block-container"]    { padding-top:1.5rem !important; }
h1,h2,h3,h4 { font-family:'Plus Jakarta Sans',sans-serif !important; color:var(--blue-800) !important; }
h1 { font-size:1.65rem !important; font-weight:800 !important; letter-spacing:-0.3px; }
h2 { font-size:1.25rem !important; font-weight:700 !important; }
h3 { font-size:1.05rem !important; font-weight:700 !important; }
.app-header {
    background:linear-gradient(135deg,var(--navy) 0%,var(--blue-800) 55%,var(--blue-600) 100%);
    border-radius:var(--radius); padding:20px 30px; margin-bottom:24px;
    display:flex; align-items:center; gap:18px;
    box-shadow:0 6px 30px rgba(13,47,82,0.28);
}
.app-header-icon {
    background:rgba(255,255,255,0.14); border-radius:12px; width:56px; height:56px;
    display:flex; align-items:center; justify-content:center; font-size:26px; flex-shrink:0;
}
.app-header h1 {
    color:#ffffff !important; margin:0 0 3px !important; font-size:1.45rem !important;
    font-weight:800 !important; letter-spacing:-0.2px !important;
    text-shadow:0 1px 4px rgba(0,0,0,.22) !important;
}
.app-header p { color:rgba(255,255,255,.68) !important; margin:0 !important; font-size:.83rem !important; }
[data-testid="stSidebar"] {
    background:linear-gradient(180deg,var(--navy) 0%,#1a4570 40%,var(--blue-800) 100%) !important;
    border-right:none !important; box-shadow:4px 0 24px rgba(13,47,82,0.28) !important;
}
[data-testid="stSidebar"] * { color:#e8f0fa !important; }
[data-testid="stSidebar"] p,[data-testid="stSidebar"] span,[data-testid="stSidebar"] div
    { font-family:'DM Sans',sans-serif !important; font-size:13.5px; }
[data-testid="stSidebar"] strong
    { color:#ffffff !important; font-family:'Plus Jakarta Sans',sans-serif !important;
      font-size:14px; font-weight:700; }
[data-testid="stSidebar"] hr { border-color:rgba(255,255,255,0.15) !important; margin:12px 0 !important; }
[data-testid="stSidebar"] [data-testid="stSelectbox"]>div>div {
    background:rgba(255,255,255,0.10) !important; border:1px solid rgba(255,255,255,0.22) !important;
    border-radius:9px !important; color:#ffffff !important; font-weight:600;
}
[data-testid="stSidebar"] [data-testid="stSelectbox"]>div>div:hover
    { background:rgba(255,255,255,0.18) !important; border-color:rgba(255,255,255,0.40) !important; }
[data-testid="stSidebar"] label
    { font-size:.75rem !important; font-weight:600 !important; letter-spacing:.8px !important;
      text-transform:uppercase !important; color:rgba(255,255,255,0.45) !important; }
[data-testid="stSidebar"] .stButton>button {
    background:rgba(239,68,68,0.18) !important; border:1px solid rgba(239,68,68,0.38) !important;
    color:#fca5a5 !important; border-radius:9px !important; font-weight:600 !important;
    height:2.6em !important; transition:all 0.2s ease !important;
}
[data-testid="stSidebar"] .stButton>button:hover
    { background:rgba(239,68,68,0.32) !important; border-color:#ef4444 !important;
      color:#ffffff !important; transform:translateY(-1px) !important; }
[data-testid="stMetric"] {
    background:var(--surface) !important; border-radius:var(--radius) !important;
    padding:18px 22px !important; box-shadow:var(--shadow) !important;
    border-left:4px solid var(--blue-400) !important;
    transition:transform 0.2s ease,box-shadow 0.2s ease !important;
}
[data-testid="stMetric"]:hover { transform:translateY(-3px) !important; box-shadow:var(--shadow-lg) !important; }
[data-testid="stMetricLabel"]
    { font-family:'DM Sans',sans-serif !important; font-size:11.5px !important;
      font-weight:500 !important; color:var(--muted) !important;
      text-transform:uppercase !important; letter-spacing:.7px !important; }
[data-testid="stMetricValue"]
    { font-family:'Plus Jakarta Sans',sans-serif !important; font-size:2rem !important;
      font-weight:800 !important; color:var(--blue-800) !important; }
.stButton>button {
    background:linear-gradient(135deg,var(--blue-600) 0%,var(--blue-400) 100%) !important;
    color:#ffffff !important; border:none !important; border-radius:10px !important;
    padding:0 1.4em !important; font-family:'Plus Jakarta Sans',sans-serif !important;
    font-size:14px !important; font-weight:600 !important; height:2.85em !important;
    box-shadow:0 2px 12px rgba(74,144,226,0.30) !important; transition:all 0.22s ease !important;
}
.stButton>button:hover {
    background:linear-gradient(135deg,var(--blue-800) 0%,var(--blue-600) 100%) !important;
    box-shadow:0 6px 20px rgba(31,78,121,0.35) !important; transform:translateY(-2px) !important;
}
.stButton>button:active { transform:translateY(0) !important; }
[data-testid="stDownloadButton"]>button {
    background:linear-gradient(135deg,#059669 0%,#10b981 100%) !important;
    box-shadow:0 2px 12px rgba(16,185,129,0.28) !important;
}
[data-testid="stDownloadButton"]>button:hover {
    background:linear-gradient(135deg,#047857 0%,#059669 100%) !important;
    box-shadow:0 6px 20px rgba(5,150,105,0.35) !important;
}
[data-testid="stTabs"] [data-baseweb="tab-list"] {
    background:var(--surface) !important; border-radius:10px !important;
    border:1px solid var(--border) !important; padding:4px !important; gap:2px !important;
    box-shadow:0 2px 8px rgba(31,78,121,0.06) !important;
}
[data-testid="stTabs"] [data-baseweb="tab"] {
    border-radius:8px !important; font-family:'Plus Jakarta Sans',sans-serif !important;
    font-size:13.5px !important; font-weight:600 !important; color:var(--muted) !important;
    padding:8px 18px !important; transition:all 0.2s ease !important;
}
[data-testid="stTabs"] [aria-selected="true"] {
    background:linear-gradient(135deg,var(--blue-800) 0%,var(--blue-600) 100%) !important;
    color:#ffffff !important; box-shadow:0 2px 8px rgba(31,78,121,0.30) !important;
}
[data-testid="stSelectbox"]>div>div,
[data-testid="stTextInput"] input,
[data-testid="stNumberInput"] input {
    border-radius:9px !important; border:1.5px solid var(--border) !important;
    font-family:'DM Sans',sans-serif !important; font-size:14px !important;
    transition:border-color .2s,box-shadow .2s !important; background:var(--surface) !important;
}
[data-testid="stTextInput"] input:focus,[data-testid="stNumberInput"] input:focus
    { border-color:var(--blue-400) !important;
      box-shadow:0 0 0 3px rgba(74,144,226,0.15) !important; outline:none !important; }
[data-testid="stTextInput"] label,[data-testid="stSelectbox"] label,
[data-testid="stNumberInput"] label {
    font-family:'Plus Jakarta Sans',sans-serif !important; font-size:12.5px !important;
    font-weight:600 !important; color:var(--blue-800) !important;
}
[data-testid="stDataFrame"]
    { border-radius:var(--radius) !important; overflow:hidden !important;
      border:1px solid var(--border) !important; box-shadow:var(--shadow) !important; }
[data-testid="stDataFrame"] td
    { white-space:pre-line !important; text-align:center !important;
      font-size:13px !important; padding:10px 14px !important;
      font-family:'DM Sans',sans-serif !important; }
[data-testid="stDataFrame"] th {
    background:var(--blue-800) !important; color:#ffffff !important;
    font-family:'Plus Jakarta Sans',sans-serif !important;
    font-weight:700 !important; font-size:12px !important;
    text-transform:uppercase !important; letter-spacing:.5px !important;
}
[data-testid="stExpander"]
    { border-radius:10px !important; border:1px solid var(--border) !important;
      box-shadow:0 2px 8px rgba(31,78,121,0.06) !important;
      margin-bottom:8px !important; overflow:hidden !important; }
[data-testid="stExpander"] summary
    { font-family:'Plus Jakarta Sans',sans-serif !important; font-weight:600 !important;
      color:var(--blue-800) !important; padding:12px 16px !important; }
[data-testid="stInfo"]    { border-radius:10px !important; border-left:4px solid var(--blue-400) !important; }
[data-testid="stWarning"] { border-radius:10px !important; border-left:4px solid var(--accent) !important; }
[data-testid="stSuccess"] { border-radius:10px !important; border-left:4px solid var(--green) !important; }
[data-testid="stError"]   { border-radius:10px !important; border-left:4px solid var(--red) !important; }
[data-testid="stSlider"] [role="slider"]
    { background:var(--blue-400) !important; border:2px solid var(--blue-800) !important; }
[data-testid="stProgressBar"]>div>div
    { background:linear-gradient(90deg,var(--blue-800),var(--blue-400)) !important; border-radius:999px !important; }
[data-testid="stProgressBar"]>div
    { background:var(--blue-100) !important; border-radius:999px !important; }
hr { border-color:var(--border) !important; margin:18px 0 !important; }
::-webkit-scrollbar { width:6px; height:6px; }
::-webkit-scrollbar-track { background:#f1f5f9; border-radius:999px; }
::-webkit-scrollbar-thumb { background:#94a3b8; border-radius:999px; }
::-webkit-scrollbar-thumb:hover { background:var(--blue-400); }
.page-header {
    background:var(--surface); border-radius:var(--radius); padding:16px 24px;
    margin-bottom:18px; border-left:5px solid var(--blue-400); box-shadow:var(--shadow);
}
.page-header h2 { margin:0 !important; font-size:1.2rem !important; }
.page-header p  { margin:4px 0 0; color:var(--muted); font-size:.85rem; }
.ui-card {
    background:var(--surface); border-radius:var(--radius); padding:22px 26px;
    box-shadow:var(--shadow); margin-bottom:16px; border:1px solid var(--border);
}
.section-label {
    font-family:'Plus Jakarta Sans',sans-serif; font-size:.7rem; font-weight:700;
    letter-spacing:1px; text-transform:uppercase; color:var(--muted); margin-bottom:6px;
}
.dl-card {
    background:var(--surface); border-radius:14px; padding:28px 32px;
    box-shadow:var(--shadow-lg); border:1px solid var(--border); margin-bottom:20px;
}
.dl-card h3 { margin:0 0 6px !important; font-size:1.05rem !important; }
.dl-card p  { color:var(--muted); font-size:.875rem; margin:0 0 18px; }

@media (max-width: 768px) {
    [data-testid="block-container"] {
        padding-left: 0.6rem !important;
        padding-right: 0.6rem !important;
        padding-top: 0.8rem !important;
    }
    .app-header {
        flex-direction: column !important;
        align-items: flex-start !important;
        gap: 10px !important;
        padding: 14px 16px !important;
        margin-bottom: 14px !important;
        border-radius: 10px !important;
    }
    .app-header-icon {
        width: 42px !important;
        height: 42px !important;
        font-size: 20px !important;
    }
    .app-header h1 { font-size: 1.05rem !important; }
    .app-header p  { font-size: 0.75rem !important; }
    h1 { font-size: 1.2rem !important; }
    h2 { font-size: 1.05rem !important; }
    h3 { font-size: 0.95rem !important; }
    .stButton > button {
        width: 100% !important;
        font-size: 13px !important;
        height: 3em !important;
    }
    [data-testid="stDownloadButton"] > button {
        width: 100% !important;
        font-size: 13px !important;
    }
    [data-testid="stDataFrame"] {
        overflow-x: auto !important;
        -webkit-overflow-scrolling: touch !important;
        max-width: 100% !important;
    }
    [data-testid="stDataFrame"] > div { overflow-x: auto !important; }
    [data-testid="stDataEditor"] {
        overflow-x: auto !important;
        -webkit-overflow-scrolling: touch !important;
        max-width: 100% !important;
    }
    [data-testid="stDataFrame"] td {
        font-size: 11px !important;
        padding: 6px 8px !important;
    }
    [data-testid="stDataFrame"] th { font-size: 10px !important; }
    [data-testid="stMetric"] { padding: 12px 14px !important; }
    [data-testid="stMetricValue"] { font-size: 1.4rem !important; }
    [data-testid="stMetricLabel"] { font-size: 10px !important; }
    [data-testid="stTabs"] [data-baseweb="tab-list"] {
        overflow-x: auto !important;
        flex-wrap: nowrap !important;
        -webkit-overflow-scrolling: touch !important;
    }
    [data-testid="stTabs"] [data-baseweb="tab"] {
        font-size: 11.5px !important;
        padding: 7px 10px !important;
        white-space: nowrap !important;
    }
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] span,
    [data-testid="stSidebar"] div { font-size: 13px !important; }
    .ui-card { padding: 14px 14px !important; }
    .dl-card { padding: 16px 14px !important; }
    div[style*="border-radius:12px"][style*="padding:16px 24px"] {
        padding: 12px 14px !important;
    }
    [data-testid="stSelectbox"] > div > div,
    [data-testid="stTextInput"] input,
    [data-testid="stNumberInput"] input { font-size: 13px !important; }
    [data-testid="stSlider"] label { font-size: 12px !important; }
    [data-testid="stExpander"] summary {
        font-size: 13px !important;
        padding: 10px 12px !important;
    }
    [data-testid="stProgressBar"] + div { font-size: 12px !important; }
    span[style*="display:inline-block"] {
        font-size: 12px !important;
        padding: 4px 10px !important;
        margin: 2px !important;
    }
    [data-testid="column"] {
        padding-left: 4px !important;
        padding-right: 4px !important;
    }
    [data-testid="stInfo"],
    [data-testid="stWarning"],
    [data-testid="stSuccess"],
    [data-testid="stError"] {
        font-size: 13px !important;
        padding: 10px 12px !important;
    }
    [data-testid="stRadio"] label { font-size: 13px !important; }
}
</style>""", unsafe_allow_html=True)

# App header
st.markdown("""
<div class="app-header">
    <div class="app-header-icon">📚</div>
    <div>
        <h1>School Timetable Scheduler Pro</h1>
        <p>DHACSS Phase IV Campus &nbsp;·&nbsp; Academic Intelligence System</p>
    </div>
</div>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# VALIDATION
# ══════════════════════════════════════════════════════════════════════════════

def validate_subject_weekly(section):
    issues, actual = [], {}
    for day in DAYS:
        for p in get_periods(day):
            s = slot(section, day, p)["subject"]
            if s: actual[s] = actual.get(s, 0) + 1
    for subj, req in st.session_state.subject_config.get(section, {}).items():
        got = actual.get(subj, 0)
        if got != req:
            issues.append(f"{subj}: needs {req} periods, has {got}")
    return issues

def validate_teacher_clashes():
    issues = []
    for day in DAYS:
        for p in get_periods(day):
            seen = {}
            for sec in st.session_state.timetable:
                t = slot(sec, day, p)["teacher"]
                if t: seen.setdefault(t, []).append(sec)
            for t, secs in seen.items():
                if len(secs) > 1:
                    issues.append(f"Clash: {t} in {secs} at {day} {p}")
    return issues

def validate_class_teacher_presence():
    issues = []
    for sec, ct in st.session_state.class_teachers.items():
        if sec not in st.session_state.timetable: continue
        found = any(slot(sec, day, p)["teacher"] == ct
                    for day in DAYS if day in st.session_state.timetable[sec]
                    for p in get_periods(day) if p in st.session_state.timetable[sec][day])
        if not found:
            issues.append(f"Class teacher {ct} has no period in {sec}")
    return issues

def validate_teacher_distribution():
    issues = []
    for t in st.session_state.teachers:
        loads = [sum(1 for sec in st.session_state.timetable for p in get_periods(day)
                     if slot(sec, day, p)["teacher"] == t) for day in DAYS]
        if max(loads) - min(loads) >= 2:
            issues.append(f"{t} daily load imbalanced: {loads}")
    return issues

def validate_friday_load():
    issues = []
    for t in st.session_state.teachers:
        n = sum(1 for sec in st.session_state.timetable for p in get_periods("Friday")
                if slot(sec, "Friday", p)["teacher"] == t)
        if n > 5: issues.append(f"{t}: heavy Friday ({n} periods)")
        elif n == 5: issues.append(f"{t}: Friday 5 periods (acceptable)")
    return issues

def validate_maths_consecutive():
    issues = []
    for sec in st.session_state.timetable:
        for day in DAYS:
            mps = [p for p in get_periods(day) if "Math" in (slot(sec,day,p)["subject"] or "")]
            if len(mps) == 2:
                ps = get_periods(day)
                if abs(ps.index(mps[0]) - ps.index(mps[1])) != 1:
                    issues.append(f"{sec}: Maths not consecutive on {day}")
    return issues

def validate_teacher_max_load(max_p=25):
    return [f"{t} exceeds {max_p} periods (has {count_teacher_periods(t)})"
            for t in st.session_state.teachers if count_teacher_periods(t) > max_p]

def validate_no_three_consecutive():
    issues = []
    for sec in st.session_state.timetable:
        for day in DAYS:
            streak, prev_t = 0, None
            for p in get_periods(day):
                if p == "Lunch": streak, prev_t = 0, None; continue
                t = slot(sec, day, p)["teacher"]
                if t and t == prev_t:
                    streak += 1
                    if streak >= 3: issues.append(f"{t}: 3+ consecutive in {sec} on {day}")
                else:
                    streak, prev_t = 1, t
    return issues

# ══════════════════════════════════════════════════════════════════════════════
# TIMETABLE QUERY HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def count_teacher_periods(teacher):
    return sum(1 for sec in st.session_state.timetable for day in DAYS
               for p in get_periods(day) if slot(sec, day, p)["teacher"] == teacher)

def teacher_busy(teacher, day, period):
    t = clean(teacher)
    return any(clean(slot(sec, day, period)["teacher"]) == t
               for sec in st.session_state.timetable)

def teacher_daily_load(teacher, day):
    t = clean(teacher)
    return sum(1 for sec in st.session_state.timetable for p in get_periods(day)
               if clean(slot(sec, day, p)["teacher"]) == t)

def subject_count_in_day(section, subject, day):
    return sum(1 for p in get_periods(day)
               if slot(section, day, p)["subject"].upper() == subject.upper())

def subject_count_total(section, subject):
    return sum(subject_count_in_day(section, subject, day) for day in DAYS)

def quota_remaining(section, subject):
    key = (section, subject)
    if key in subject_remaining:
        return subject_remaining[key]
    quota = st.session_state.subject_config.get(section, {}).get(subject, 0)
    return max(0, quota - subject_count_total(section, subject))

def math_double_day(section):
    for day in DAYS:
        ps = [p for p in get_periods(day) if p != "Lunch"]
        for i in range(len(ps)-1):
            if is_math(slot(section,day,ps[i])["subject"]) and \
               is_math(slot(section,day,ps[i+1])["subject"]):
                return day
    return None

def get_last_teaching_period(day):
    return [p for p in get_periods(day) if p != "Lunch"][-1]

# ══════════════════════════════════════════════════════════════════════════════
# SUBJECT CATEGORY HELPERS
# ══════════════════════════════════════════════════════════════════════════════

DAILY_SINGLE_KEYWORDS = [
    "SINDHI","ISLAMIAT","ENGLISH","URDU","GENERAL SCIENCE",
    "ARABIC","GEOGRAPHY","HISTORY","SOCIAL STUDIES",
]
IX_X_DOUBLE_SUBJECTS = [
    "PHYSICS","CHEMISTRY","COMPUTER IX-X","COMPUTER SCIENCE IX-X",
]
MATH_KEYWORD  = "MATH"
GAMES_KEYWORD = "GAMES"
_DS_SORTED = sorted(DAILY_SINGLE_KEYWORDS, key=len, reverse=True)

def is_daily_single(subject):
    s = subject.strip().upper()
    return any(kw in s for kw in _DS_SORTED)

def is_math(subject):        return MATH_KEYWORD  in subject.upper()
def is_games(subject):       return GAMES_KEYWORD in subject.upper()
def is_ix_x_double(subject):
    s = subject.upper()
    return any(t in s for t in IX_X_DOUBLE_SUBJECTS)
def is_double_allowed(subject):
    return is_math(subject) or is_ix_x_double(subject)

# ══════════════════════════════════════════════════════════════════════════════
# CONSTRAINT ENGINE
# ══════════════════════════════════════════════════════════════════════════════

def _init_teacher(t_key):
    if t_key not in teacher_day_load:
        teacher_day_load[t_key] = {d:0 for d in DAYS}
        teacher_timeline[t_key] = {d:[0]*len(get_periods(d)) for d in DAYS}

def teacher_consecutive_streak(teacher, day, proposed_idx):
    periods  = get_periods(day)
    occupied = []
    for i, p in enumerate(periods):
        if p == "Lunch": occupied.append(-1); continue
        if i == proposed_idx: occupied.append(1); continue
        occupied.append(1 if any(slot(sec,day,p)["teacher"]==teacher
                                  for sec in st.session_state.timetable) else 0)
    max_s = s = 0
    for v in occupied:
        if v == -1: s = 0
        elif v == 1: s += 1; max_s = max(max_s, s)
        else: s = 0
    return max_s

def is_library(subject):
    return "LIBRARY" in subject.upper()

CORE_SUBJECTS = ["MATH", "ENGLISH", "URDU", "GENERAL SCIENCE"]

def is_core(subject):
    return any(c in subject.upper() for c in CORE_SUBJECTS)

def can_assign(section, subject, teacher, day, period):
    periods = get_periods(day)
    idx     = periods.index(period)
    if period == "Lunch": return False

    if teacher_busy(teacher, day, period): return False

    t_key = clean(teacher)
    _init_teacher(t_key)
    if teacher_day_load[t_key][day] >= (5 if day == "Friday" else 6): return False

    if teacher_consecutive_streak(t_key, day, idx) >= 4: return False

    slots_list = [p for p in periods if p != "Lunch"]
    si         = slots_list.index(period) if period in slots_list else -1
    prev_p     = slots_list[si - 1] if si > 0 else None
    next_p     = slots_list[si + 1] if si < len(slots_list) - 1 else None
    lunch_i    = periods.index("Lunch") if "Lunch" in periods else -1
    period_i   = periods.index(period)

    def real_adj(other_p):
        if other_p is None: return False
        oi  = periods.index(other_p)
        lo, hi = min(period_i, oi), max(period_i, oi)
        if lunch_i != -1 and lo < lunch_i < hi: return False
        return abs(period_i - oi) == 1

    prev_s = slot(section, day, prev_p)["subject"] if prev_p and real_adj(prev_p) else ""
    next_s = slot(section, day, next_p)["subject"] if next_p and real_adj(next_p) else ""
    su     = subject.upper()

    if is_daily_single(subject):
        if subject_count_in_day(section, subject, day) >= 1: return False
        if prev_s.upper() == su or next_s.upper() == su:     return False

    if not is_double_allowed(subject):
        if subject_count_in_day(section, subject, day) >= 1: return False

    if not is_double_allowed(subject):
        if prev_s.upper() == su or next_s.upper() == su: return False

    if is_math(subject):
        today_count  = subject_count_in_day(section, subject, day)
        would_be_double = (prev_s.upper() == su or next_s.upper() == su)
        if would_be_double:
            ex = math_double_day(section)
            if ex is not None and ex != day: return False
            if today_count >= 2: return False
        else:
            ex = math_double_day(section)
            if today_count >= 1:
                return False

    if is_ix_x_double(subject):
        would_be_double = (prev_s.upper() == su or next_s.upper() == su)
        if would_be_double and double_used.get((section, subject), False):
            return False
        if not would_be_double and subject_count_in_day(section, subject, day) >= 1:
            return False

    if is_games(subject) or is_library(subject):
        teaching = [p for p in get_periods(day) if p != "Lunch"]
        if period in (teaching[0], teaching[-1]): return False

    return True

def apply_assignment(section, subject, teacher, day, period):
    t_key = clean(teacher)
    st.session_state.timetable[section][day][period] = {"subject":subject,"teacher":t_key}
    idx = get_periods(day).index(period)
    _init_teacher(t_key)
    teacher_day_load[t_key][day]     += 1
    teacher_timeline[t_key][day][idx] = 1
    streak = teacher_consecutive_streak(t_key, day, idx)
    teacher_three_streak_count.setdefault(t_key, 0)
    if streak == 3:
        teacher_three_streak_count[t_key] += 1
    key = (section, subject)
    if key in subject_remaining:
        subject_remaining[key] = max(0, subject_remaining[key]-1)
    if is_ix_x_double(subject):
        ps = [p for p in get_periods(day) if p != "Lunch"]
        for i, p in enumerate(ps):
            if p == period:
                if (i>0 and slot(section,day,ps[i-1])["subject"].upper()==subject.upper()) or \
                   (i<len(ps)-1 and slot(section,day,ps[i+1])["subject"].upper()==subject.upper()):
                    double_used[(section,subject)] = True
                break

def undo_assignment(section, day, period):
    cell    = st.session_state.timetable[section][day][period]
    subject = cell.get("subject","")
    teacher = cell.get("teacher","")
    if not subject: return
    key = (section, subject)
    if key in subject_remaining: subject_remaining[key] += 1
    t_key = clean(teacher) if teacher else ""
    if t_key and t_key in teacher_day_load:
        teacher_day_load[t_key][day] = max(0, teacher_day_load[t_key][day]-1)
    if t_key and t_key in teacher_timeline:
        teacher_timeline[t_key][day][get_periods(day).index(period)] = 0
    st.session_state.timetable[section][day][period] = {"subject":"","teacher":""}
    if is_ix_x_double(subject):
        still = any(
            slot(section,d,ps[i])["subject"].upper()==subject.upper() and
            slot(section,d,ps[i+1])["subject"].upper()==subject.upper()
            for d in DAYS
            for ps in [[p for p in get_periods(d) if p!="Lunch"]]
            for i in range(len(ps)-1)
        )
        if not still: double_used.pop((section,subject), None)

# ══════════════════════════════════════════════════════════════════════════════
# SCHEDULER HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _find_teacher(section, subject):
    for t, sec_data in st.session_state.teacher_assignment.items():
        if section in sec_data and subject in sec_data[section]:
            return clean(t)
    return None

def _is_ix_x_section(section):
    return bool(re.search(r'(?<![A-Z])(?:IX|X|9|10)(?![A-Z0-9])', section.upper()))

def create_empty_timetable():
    teacher_day_load.clear(); teacher_timeline.clear()
    double_used.clear();      subject_remaining.clear()
    teacher_three_streak_count.clear()
    for t in st.session_state.teachers:
        teacher_three_streak_count[clean(t)] = 0
    for sec, sm in st.session_state.subject_config.items():
        for subj, quota in sm.items():
            subject_remaining[(sec, subj)] = quota
    for t in st.session_state.teachers:
        _init_teacher(clean(t))
    return {sec:{day:{p:{"subject":"","teacher":""}
                      for p in get_periods(day)} for day in DAYS}
            for sec in st.session_state.sections}

def assign_class_teacher_priority():
    for sec, ct in st.session_state.class_teachers.items():
        subjects = st.session_state.teacher_assignment.get(ct,{}).get(sec,[])
        if not subjects: continue
        assigned = 0
        for day in DAYS:
            if "P1" not in get_periods(day): continue
            for subj in subjects:
                if quota_remaining(sec,subj)<=0: continue
                if slot(sec,day,"P1")["subject"]=="" and can_assign(sec,subj,ct,day,"P1"):
                    apply_assignment(sec,subj,ct,day,"P1")
                    assigned += 1; break
            if assigned >= 4: break

def assign_daily_singles():
    for sec in st.session_state.subject_config:
        for subj in list(st.session_state.subject_config[sec]):
            if not is_daily_single(subj): continue
            teacher = _find_teacher(sec, subj)
            if not teacher: continue
            days_to_try = DAYS.copy()
            random.shuffle(days_to_try)
            for day in days_to_try:
                if quota_remaining(sec, subj) <= 0: break
                if subject_count_in_day(sec, subj, day) >= 1: continue
                cands = [p for p in get_periods(day) if p != "Lunch"]
                random.shuffle(cands)
                for p in cands:
                    if quota_remaining(sec, subj) <= 0: break
                    if slot(sec, day, p)["subject"] != "": continue
                    if can_assign(sec, subj, teacher, day, p):
                        apply_assignment(sec, subj, teacher, day, p)
                        break

def assign_math():
    for sec in st.session_state.subject_config:
        for subj in st.session_state.subject_config[sec]:
            if not is_math(subj): continue
            teacher = _find_teacher(sec, subj)
            if not teacher: continue
            double_day = None
            if quota_remaining(sec, subj) >= 2:
                days_shuffled = DAYS.copy(); random.shuffle(days_shuffled)
                for day in days_shuffled:
                    if quota_remaining(sec, subj) < 2: break
                    ps = [p for p in get_periods(day) if p != "Lunch"]
                    for i in range(len(ps) - 1):
                        p1, p2 = ps[i], ps[i + 1]
                        if (slot(sec, day, p1)["subject"] == ""
                                and slot(sec, day, p2)["subject"] == ""
                                and can_assign(sec, subj, teacher, day, p1)
                                and can_assign(sec, subj, teacher, day, p2)):
                            apply_assignment(sec, subj, teacher, day, p1)
                            apply_assignment(sec, subj, teacher, day, p2)
                            double_day = day; break
                    if double_day: break
            for day in DAYS:
                if quota_remaining(sec, subj) <= 0: break
                if day == double_day: continue
                if subject_count_in_day(sec, subj, day) >= 1: continue
                cands = [p for p in get_periods(day) if p != "Lunch"]
                random.shuffle(cands)
                for p in cands:
                    if quota_remaining(sec, subj) <= 0: break
                    if slot(sec, day, p)["subject"] == "" and can_assign(sec, subj, teacher, day, p):
                        apply_assignment(sec, subj, teacher, day, p); break

def assign_ix_x_doubles():
    for sec in st.session_state.subject_config:
        if not _is_ix_x_section(sec): continue
        for subj in list(st.session_state.subject_config[sec]):
            if not is_ix_x_double(subj): continue
            teacher = _find_teacher(sec, subj)
            if not teacher: continue
            double_day = None
            if not double_used.get((sec, subj), False) and quota_remaining(sec, subj) >= 2:
                days_shuffled = DAYS.copy(); random.shuffle(days_shuffled)
                for day in days_shuffled:
                    if quota_remaining(sec, subj) < 2: break
                    ps = [p for p in get_periods(day) if p != "Lunch"]
                    for i in range(len(ps) - 1):
                        p1, p2 = ps[i], ps[i + 1]
                        if (slot(sec, day, p1)["subject"] == ""
                                and slot(sec, day, p2)["subject"] == ""
                                and can_assign(sec, subj, teacher, day, p1)
                                and can_assign(sec, subj, teacher, day, p2)):
                            apply_assignment(sec, subj, teacher, day, p1)
                            apply_assignment(sec, subj, teacher, day, p2)
                            double_day = day; break
                    if double_day: break
            for day in DAYS:
                if quota_remaining(sec, subj) <= 0: break
                if day == double_day: continue
                if subject_count_in_day(sec, subj, day) >= 1: continue
                cands = [p for p in get_periods(day) if p != "Lunch"]
                random.shuffle(cands)
                for p in cands:
                    if quota_remaining(sec, subj) <= 0: break
                    if slot(sec, day, p)["subject"] == "" and can_assign(sec, subj, teacher, day, p):
                        apply_assignment(sec, subj, teacher, day, p); break

def try_swap(section, subject, teacher):
    for day in DAYS:
        for period in get_periods(day):
            if period=="Lunch": continue
            curr = st.session_state.timetable[section][day][period]
            if not curr["subject"]: continue
            os2, ot = curr["subject"], curr["teacher"]
            undo_assignment(section, day, period)
            if not teacher_busy(teacher,day,period) and can_assign(section,subject,teacher,day,period):
                for d2 in DAYS:
                    for p2 in get_periods(d2):
                        if p2=="Lunch" or (d2,p2)==(day,period): continue
                        if st.session_state.timetable[section][d2][p2]["subject"]: continue
                        if not teacher_busy(ot,d2,p2) and can_assign(section,os2,ot,d2,p2):
                            apply_assignment(section,subject,teacher,day,period)
                            apply_assignment(section,os2,ot,d2,p2)
                            return True
            apply_assignment(section,os2,ot,day,period)
    return False

def try_displace(section, subject, teacher):
    for day in DAYS:
        for period in get_periods(day):
            if period == "Lunch": continue
            curr = st.session_state.timetable[section][day][period]
            os2, ot = curr.get("subject",""), curr.get("teacher","")
            if not os2 or quota_remaining(section, os2) > 0: continue
            undo_assignment(section, day, period)
            if (not teacher_busy(teacher, day, period)
                    and can_assign(section, subject, teacher, day, period)):
                apply_assignment(section, subject, teacher, day, period)
                return True
            apply_assignment(section, os2, ot, day, period)
    return False

def basic_auto_fill():
    secs = list(st.session_state.subject_config.keys()); random.shuffle(secs)
    for sec in secs:
        items = sorted(st.session_state.subject_config[sec].items(),
                       key=lambda kv: quota_remaining(sec, kv[0]), reverse=True)
        for subj, _ in items:
            teacher = _find_teacher(sec, subj)
            if not teacher: continue
            while quota_remaining(sec, subj) > 0:
                zero_days = [d for d in DAYS if subject_count_in_day(sec, subj, d) == 0]
                one_days  = [d for d in DAYS if subject_count_in_day(sec, subj, d) == 1
                             and is_double_allowed(subj)]
                random.shuffle(zero_days); random.shuffle(one_days)
                ordered_days = zero_days + one_days
                def day_cap(d):
                    if is_math(subj):
                        return 2 if d == math_double_day(sec) else 1
                    if is_double_allowed(subj): return 2
                    return 1
                valid = []
                for day in ordered_days:
                    if subject_count_in_day(sec, subj, day) >= day_cap(day): continue
                    ps = get_periods(day).copy(); random.shuffle(ps)
                    for p in ps:
                        if p == "Lunch" or slot(sec, day, p)["subject"]: continue
                        if teacher_busy(teacher, day, p): continue
                        if teacher_daily_load(teacher, day) >= (5 if day == "Friday" else 6): continue
                        if can_assign(sec, subj, teacher, day, p):
                            valid.append((day, p))
                if not valid:
                    if try_swap(sec, subj, teacher): continue
                    if try_displace(sec, subj, teacher): continue
                    break
                day, p = random.choice(valid)
                apply_assignment(sec, subj, teacher, day, p)

def fill_under_quota_subjects():
    for sec in st.session_state.subject_config:
        for subj in sorted(st.session_state.subject_config[sec],
                           key=lambda s: quota_remaining(sec, s), reverse=True):
            teacher = _find_teacher(sec, subj)
            if not teacher: continue
            def strict_day_cap(d):
                if is_math(subj): return 2
                if is_double_allowed(subj): return 2
                return 1
            for relax in (False, True):
                if quota_remaining(sec, subj) <= 0: break
                days_shuffled = DAYS.copy(); random.shuffle(days_shuffled)
                for day in days_shuffled:
                    if quota_remaining(sec, subj) <= 0: break
                    cap = strict_day_cap(day) if relax else 1
                    if is_double_allowed(subj) and not relax: cap = 2
                    if not is_double_allowed(subj): cap = 1
                    if subject_count_in_day(sec, subj, day) >= cap: continue
                    ps = [p for p in get_periods(day) if p != "Lunch"]
                    random.shuffle(ps)
                    for p in ps:
                        if quota_remaining(sec, subj) <= 0: break
                        if slot(sec, day, p)["subject"]: continue
                        if can_assign(sec, subj, teacher, day, p):
                            apply_assignment(sec, subj, teacher, day, p); break
            while quota_remaining(sec, subj) > 0:
                if not try_swap(sec, subj, teacher): break
            while quota_remaining(sec, subj) > 0:
                if not try_displace(sec, subj, teacher): break

def emergency_backfill():
    for sec in st.session_state.timetable:
        for day in DAYS:
            for period in get_periods(day):
                if period=="Lunch" or slot(sec,day,period)["subject"]: continue
                for subj,_ in sorted(st.session_state.subject_config.get(sec,{}).items(),
                                     key=lambda kv: quota_remaining(sec,kv[0]), reverse=True):
                    if quota_remaining(sec,subj)<=0: continue
                    teacher=_find_teacher(sec,subj)
                    if teacher and can_assign(sec,subj,teacher,day,period):
                        apply_assignment(sec,subj,teacher,day,period); break
    for sec in st.session_state.subject_config:
        for subj in sorted([s for s in st.session_state.subject_config[sec]
                            if quota_remaining(sec,s)>0],
                           key=lambda s: quota_remaining(sec,s), reverse=True):
            teacher=_find_teacher(sec,subj)
            if not teacher: continue
            stalled=False
            while quota_remaining(sec,subj)>0 and not stalled:
                placed=False
                for day in DAYS:
                    if placed: break
                    for period in get_periods(day):
                        if period=="Lunch" or slot(sec,day,period)["subject"]: continue
                        if can_assign(sec,subj,teacher,day,period):
                            apply_assignment(sec,subj,teacher,day,period); placed=True; break
                if placed: continue
                if try_swap(sec,subj,teacher): continue
                if try_displace(sec,subj,teacher): continue
                stalled=True

def ensure_class_teacher_presence():
    for sec, ct in st.session_state.class_teachers.items():
        if sec not in st.session_state.timetable: continue
        if any(slot(sec, day, p)["teacher"] == ct
               for day in DAYS for p in get_periods(day) if p != "Lunch"):
            continue
        subjects = st.session_state.teacher_assignment.get(ct, {}).get(sec, [])
        if not subjects: continue
        placed = False
        for day in DAYS:
            if placed: break
            for p in get_periods(day):
                if p == "Lunch" or slot(sec, day, p)["subject"]: continue
                for subj in subjects:
                    if quota_remaining(sec, subj) <= 0: continue
                    if can_assign(sec, subj, ct, day, p):
                        apply_assignment(sec, subj, ct, day, p)
                        placed = True; break
                if placed: break
        if placed: continue
        for day in DAYS:
            if placed: break
            for p in get_periods(day):
                if p == "Lunch": continue
                cell = slot(sec, day, p)
                if cell["teacher"] == ct: break
                if cell["subject"] not in subjects: continue
                subj  = cell["subject"]
                old_t = cell["teacher"]
                if teacher_busy(ct, day, p): continue
                undo_assignment(sec, day, p)
                if can_assign(sec, subj, ct, day, p):
                    apply_assignment(sec, subj, ct, day, p)
                    placed = True; break
                else:
                    apply_assignment(sec, subj, old_t, day, p)
            if placed: break

def force_fill():
    # Pass 1
    for sec in st.session_state.timetable:
        for day in DAYS:
            for period in get_periods(day):
                if period == "Lunch" or slot(sec, day, period)["subject"]: continue
                candidates = sorted(
                    [(subj, quota_remaining(sec, subj))
                     for subj in st.session_state.subject_config.get(sec, {})
                     if quota_remaining(sec, subj) > 0],
                    key=lambda x: -x[1])
                for subj, _ in candidates:
                    t = _find_teacher(sec, subj)
                    if t and can_assign(sec, subj, t, day, period):
                        apply_assignment(sec, subj, t, day, period); break
    # Pass 2
    for sec in st.session_state.timetable:
        for day in DAYS:
            for period in get_periods(day):
                if period == "Lunch" or slot(sec, day, period)["subject"]: continue
                periods_list = get_periods(day)
                idx = periods_list.index(period)
                for subj in sorted(st.session_state.subject_config.get(sec, {}),
                                   key=lambda s: -quota_remaining(sec, s)):
                    t = _find_teacher(sec, subj)
                    if not t: continue
                    tk = clean(t)
                    _init_teacher(tk)
                    if teacher_busy(t, day, period): continue
                    if teacher_day_load[tk][day] >= (5 if day == "Friday" else 6): continue
                    if teacher_consecutive_streak(tk, day, idx) >= 4: continue
                    if is_games(subj) or is_library(subj):
                        teaching = [p for p in get_periods(day) if p != "Lunch"]
                        if period in (teaching[0], teaching[-1]): continue
                    apply_assignment(sec, subj, t, day, period); break
    # Pass 3
    for sec in st.session_state.timetable:
        for day in DAYS:
            for period in get_periods(day):
                if period == "Lunch" or slot(sec, day, period)["subject"]: continue
                placed = False
                for subj in st.session_state.subject_config.get(sec, {}):
                    if placed: break
                    t = _find_teacher(sec, subj)
                    if not t: continue
                    tk = clean(t)
                    _init_teacher(tk)
                    if teacher_busy(t, day, period): continue
                    if teacher_day_load[tk][day] >= (5 if day == "Friday" else 6): continue
                    apply_assignment(sec, subj, t, day, period)
                    placed = True
                if not placed:
                    for subj in st.session_state.subject_config.get(sec, {}):
                        t = _find_teacher(sec, subj)
                        if t and not teacher_busy(t, day, period):
                            apply_assignment(sec, subj, t, day, period); break

def calculate_fitness():
    score = 10_000
    for sec,cfg in st.session_state.subject_config.items():
        for subj in cfg: score -= quota_remaining(sec,subj)*200
    score -= sum(teacher_three_streak_count.values())*30
    score -= len(validate_no_three_consecutive())*50
    score -= len(validate_teacher_distribution())*20
    score -= len(validate_friday_load())*10
    return score

# ══════════════════════════════════════════════════════════════════════════════
# MISC HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def get_timetable_text():
    text = ""
    for sec in st.session_state.timetable:
        text += f"\nSection: {sec}\n"
        for day in DAYS:
            text += f"{day}: "
            for p in get_periods(day):
                if p == "Lunch":
                    continue
                cell = st.session_state.timetable[sec][day][p]
                if cell["teacher"]:
                    text += f"{p}-{cell['teacher']} ({cell['subject']}), "
            text += "\n"
    return text
def build_daywise_class_view(day):
    periods = get_periods(day)
    data = []

    time_row = {"Class": "Time"}
    for p in periods:
        if p == "Lunch":
            time_row["Break"] = ""
        else:
            time_row[p] = get_time(day, p)
    data.append(time_row)

    for sec in st.session_state.timetable:
        row = {"Class": sec}
        for p in periods:
            if p == "Lunch":
                row["Break"] = "BREAK"
                continue
            cell = slot(sec, day, p)
            if cell["subject"]:
                row[p] = f"{cell['subject']} ({cell['teacher']})"
            else:
                row[p] = ""
        data.append(row)

    df = pd.DataFrame(data)
    cols = ["Class"] + [p for p in periods if p != "Lunch"]
    if "Lunch" in periods:
        cols.insert(5, "Break")
    return df[cols]

def export_daywise_excel(day):
    df = build_daywise_class_view(day)
    file_name = f"{day}_Class_Timetable.xlsx"
    with pd.ExcelWriter(file_name, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, startrow=5, sheet_name=day)
        ws = writer.sheets[day]
        ws["A1"] = "DHACSS PHASE IV"
        ws["A2"] = "DAY WISE TIME TABLE (SENIOR SECTION)"
        ws["A3"] = f"DAY: {day}"
        for row in ["A1", "A2", "A3"]:
            ws[row].font = Font(bold=True, size=14)
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(df.columns))
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(df.columns))
        ws.merge_cells(start_row=3, start_column=1, end_row=3, end_column=len(df.columns))
        header_fill = PatternFill("solid", fgColor="1F4E79")
        header_row = 6
        for cell in ws[header_row]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
        time_row_index = header_row + 1
        for cell in ws[time_row_index]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill("solid", fgColor="D9E1F2")
            cell.alignment = Alignment(horizontal="center", vertical="center")
        thin = Side(style="thin")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)
        for row in ws.iter_rows(min_row=header_row+1):
            for cell in row:
                cell.border = border
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for col in ws.columns:
            max_length = max(len(str(cell.value or "")) for cell in col)
            ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_length + 5, 35)
    return file_name

def replace_teacher_everywhere(old, new):
    for sec in st.session_state.timetable:
        for day in DAYS:
            for p in get_periods(day):
                if st.session_state.timetable[sec][day][p]["teacher"]==old:
                    st.session_state.timetable[sec][day][p]["teacher"]=new
    save_all_data()

def _pdf_style(lunch_rows=None):
    cmds = [
        ("BACKGROUND",(0,0),(-1,0), colors.HexColor(BRAND_BLUE)),
        ("TEXTCOLOR",(0,0),(-1,0),  colors.white),
        ("FONTNAME",(0,0),(-1,0),   "Helvetica-Bold"),
        ("ALIGN",(0,0),(-1,-1),     "CENTER"),
        ("VALIGN",(0,0),(-1,-1),    "MIDDLE"),
        ("FONTSIZE",(0,0),(-1,-1),  9),
        ("GRID",(0,0),(-1,-1),      0.5, colors.grey),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white, colors.HexColor("#eaf0fb")]),
    ]
    if lunch_rows:
        for r in lunch_rows:
            cmds.append(("BACKGROUND",(0,r),(-1,r),colors.HexColor("#fff2cc")))
    return TableStyle(cmds)

# ══════════════════════════════════════════════════════════════════════════════
# EXPORTS
# ══════════════════════════════════════════════════════════════════════════════

def _word_landscape(doc):
    s = doc.sections[-1]
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = s.page_height, s.page_width

def export_teacher_view_word(teacher):
    doc = Document(); _word_landscape(doc)
    doc.add_heading(f"Teacher Timetable — {teacher}", 0)
    doc.add_paragraph(f"Total Weekly Periods: {count_teacher_periods(teacher)}")
    headers = ["Period","Mon-Wed Time","Mon","Tue","Wed","Thu","Fri","Fri Time"]
    table   = doc.add_table(rows=len(ALL_PERIODS)+1, cols=len(headers))
    table.style = "Table Grid"
    for i,h in enumerate(headers):
        c=table.cell(0,i); c.text=h; c.paragraphs[0].runs[0].bold=True
    for r,period in enumerate(ALL_PERIODS):
        table.cell(r+1,0).text = period
        table.cell(r+1,1).text = get_time("Monday",period)
        col=2
        for day in DAYS:
            if day=="Friday" and period not in get_periods(day):
                table.cell(r+1,col).text=""
            elif period in get_periods(day):
                val=""
                for sec in st.session_state.timetable:
                    e=slot(sec,day,period)
                    if clean(e["teacher"])==clean(teacher): val=f"{sec}\n{e['subject']}"; break
                table.cell(r+1,col).text=val
            col+=1
        table.cell(r+1,col).text=FRIDAY_TIMES.get(period,"")
    fn=f"{teacher}_timetable.docx"; doc.save(fn); return fn

def export_teacher_view_pdf(teacher):
    path=f"{teacher}_timetable.pdf"
    doc=SimpleDocTemplate(path,pagesize=landscape(A4),
                          leftMargin=20,rightMargin=20,topMargin=30,bottomMargin=30)
    styles=getSampleStyleSheet()
    elems=[Paragraph(f"<b>Teacher Timetable — {teacher}</b>",styles["Title"]),Spacer(1,8),
           Paragraph(f"<b>Total Weekly Periods:</b> {count_teacher_periods(teacher)}",styles["Normal"]),
           Spacer(1,14)]
    data=[["Period","Mon-Wed Time","Mon","Tue","Wed","Thu","Fri","Fri Time"]]
    lunch_rows=[]
    for ri,period in enumerate(ALL_PERIODS,1):
        if period=="Lunch": lunch_rows.append(ri)
        row=[period,get_time("Monday",period)]
        for day in DAYS:
            if day=="Friday" and period not in get_periods(day): row.append(""); continue
            val=""
            if period in get_periods(day):
                for sec in st.session_state.timetable:
                    e=slot(sec,day,period)
                    if clean(e["teacher"])==clean(teacher): val=f"{sec}\n{e['subject']}"; break
            row.append(val)
        row.append(FRIDAY_TIMES.get(period,"")); data.append(row)
    pw=landscape(A4)[0]-40; cw=pw/len(data[0])
    tbl=Table(data,colWidths=[cw]*len(data[0]),repeatRows=1)
    tbl.setStyle(_pdf_style(lunch_rows)); elems.append(tbl); doc.build(elems)
    return path

def export_class_timetable_pdf(section):
    path=f"{section}_timetable.pdf"
    doc=SimpleDocTemplate(path,pagesize=landscape(A4),
                          leftMargin=40,rightMargin=40,topMargin=40,bottomMargin=40)
    styles=getSampleStyleSheet()
    ct=st.session_state.class_teachers.get(section,"Not Assigned")
    elems=[Paragraph(f"<b>Class Timetable — {section}</b>",styles["Title"]),Spacer(1,8),
           Paragraph(f"<b>Class Teacher:</b> {ct}",styles["Normal"]),Spacer(1,18)]
    data=[["Period","Mon-Wed","Mon","Tue","Wed","Thu Time","Thursday","Friday","Fri Time"]]
    lunch_rows=[]
    for ri,period in enumerate(ALL_PERIODS,1):
        if period=="Lunch": lunch_rows.append(ri)
        row=[period, MON_WED_TIMES.get(period,"")]
        for day in ["Monday","Tuesday","Wednesday"]:
            c=slot(section,day,period) if period in get_periods(day) else {}
            row.append(f"{c['subject']}\n({c['teacher']})" if c.get("subject") else "")
        row.append(THURSDAY_TIMES.get(period,""))
        c=slot(section,"Thursday",period) if period in get_periods("Thursday") else {}
        row.append(f"{c['subject']}\n({c['teacher']})" if c.get("subject") else "")
        c=slot(section,"Friday",period) if period in get_periods("Friday") else {}
        row.append(f"{c['subject']}\n({c['teacher']})" if c.get("subject") else "")
        row.append(FRIDAY_TIMES.get(period,"")); data.append(row)
    cw = [50, 75, 110, 110, 110, 75, 110, 110, 75]
    tbl=Table(data,colWidths=cw,repeatRows=1)
    tbl.setStyle(_pdf_style(lunch_rows)); elems.append(tbl); doc.build(elems)
    return path

def export_class_timetable_to_word(section):
    doc=Document(); _word_landscape(doc)
    ct=st.session_state.class_teachers.get(section,"Not Assigned")
    doc.add_heading(f"Class Timetable — {section}",0)
    doc.add_paragraph(f"Class Teacher: {ct}")
    headers=["Period","Mon-Wed","Monday","Tuesday","Wednesday","Thu Time","Thursday","Friday","Fri Time"]
    table=doc.add_table(rows=len(ALL_PERIODS)+1,cols=len(headers))
    table.style="Table Grid"
    for i,h in enumerate(headers):
        c=table.cell(0,i); c.text=h; c.paragraphs[0].runs[0].bold=True
    for r,period in enumerate(ALL_PERIODS):
        table.cell(r+1,0).text=period
        table.cell(r+1,1).text=MON_WED_TIMES.get(period,"")
        col=2
        for day in ["Monday","Tuesday","Wednesday"]:
            c=slot(section,day,period) if period in get_periods(day) else {}
            table.cell(r+1,col).text=(f"{c['subject']} ({c['teacher']})" if c.get("subject") else ""); col+=1
        table.cell(r+1,col).text=THURSDAY_TIMES.get(period,""); col+=1
        c=slot(section,"Thursday",period) if period in get_periods("Thursday") else {}
        table.cell(r+1,col).text=(f"{c['subject']} ({c['teacher']})" if c.get("subject") else ""); col+=1
        c=slot(section,"Friday",period) if period in get_periods("Friday") else {}
        table.cell(r+1,col).text=(f"{c['subject']} ({c['teacher']})" if c.get("subject") else ""); col+=1
        table.cell(r+1,col).text=FRIDAY_TIMES.get(period,"")
    fn=f"{section}_timetable.docx"; doc.save(fn); return fn

def build_principal_matrix():
    if not st.session_state.timetable: return pd.DataFrame()
    rows=[]
    for day in DAYS:
        pno=0
        for period in get_periods(day):
            if period!="Lunch": pno+=1
            row={"Day":day,"P. No.":"" if period=="Lunch" else pno,"Bell Timing":get_time(day,period)}
            for t in st.session_state.teachers:
                val=""
                for sec in st.session_state.timetable:
                    if period not in st.session_state.timetable[sec].get(day,{}): continue
                    e=slot(sec,day,period)
                    if e["teacher"]==t: val=f"{sec} {e['subject']}"; break
                row[t]=val
            rows.append(row)
    df=pd.DataFrame(rows)
    df.loc[df["Day"].duplicated(),"Day"]=""
    df["P. No."]=df["P. No."].astype(str)
    cols = list(df.columns)
    fixed_cols = ["Day", "P. No.", "Bell Timing"]
    teacher_cols = [c for c in cols if c not in fixed_cols]

    def teacher_group(teacher):
        subjects = []
        for sec in st.session_state.teacher_assignment.get(teacher, {}):
            subjects += st.session_state.teacher_assignment[teacher][sec]
        for s in subjects:
            return get_group(s)
        return "OTHER"

    GROUP_ORDER = [
        "ENGLISH","URDU","ISLAMIAT_ARABIC","MATH","SCIENCE",
        "COMPUTER","SOCIAL","SINDHI","TS"
    ]
    teacher_cols = sorted(
        teacher_cols,
        key=lambda t: GROUP_ORDER.index(teacher_group(t)) if teacher_group(t) in GROUP_ORDER else 999
    )
    df = df[fixed_cols + teacher_cols]
    return df

def export_timetable_word(df):
    file="school_timetable.docx"; doc=Document(); _word_landscape(doc)
    for i,day in enumerate(DAYS):
        day_df=df.iloc[i*9:(i+1)*9]
        doc.add_heading(day,level=1)
        table=doc.add_table(rows=len(day_df)+1,cols=len(day_df.columns))
        table.style="Table Grid"
        for j,col in enumerate(day_df.columns): table.rows[0].cells[j].text=str(col)
        for r,(_,row) in enumerate(day_df.iterrows(),1):
            for ci,val in enumerate(row): table.rows[r].cells[ci].text=str(val)
        doc.add_paragraph("")
    doc.save(file); return file

def export_excel(df):
    file="School_Timetable.xlsx"
    with pd.ExcelWriter(file,engine="openpyxl") as writer:
        df.to_excel(writer,sheet_name="Timetable",index=False)
        ws=writer.sheets["Timetable"]
        ws.freeze_panes="A2"
        hf=PatternFill("solid",fgColor="1F4E79")
        for cell in ws[1]:
            cell.font=Font(bold=True,color="FFFFFF")
            cell.alignment=Alignment(horizontal="center",vertical="center")
            cell.fill=hf
        thin=Side(style="thin"); border=Border(left=thin,right=thin,top=thin,bottom=thin)
        alt=PatternFill("solid",fgColor="EAF0FB")
        lf =PatternFill("solid",fgColor="FFF2CC")
        for r,row in enumerate(ws.iter_rows(min_row=2),2):
            is_lunch = str(ws.cell(r,2).value)==""
            for cell in row:
                cell.border=border
                cell.alignment=Alignment(horizontal="center",vertical="center",wrap_text=True)
                if is_lunch:
                    cell.fill = lf
                elif r % 2 == 0:
                    cell.fill = alt
                else:
                    cell.fill = PatternFill(fill_type=None)
        for col in ws.columns:
            ml=max((len(str(c.value or "")) for c in col),default=0)
            ws.column_dimensions[get_column_letter(col[0].column)].width=min(ml+4,30)
    return file

# ══════════════════════════════════════════════════════════════════════════════
# BULK ZIP HELPER
# ══════════════════════════════════════════════════════════════════════════════

def create_zip(file_paths: list, zip_name: str = "timetable_download.zip") -> str:
    with zipfile.ZipFile(zip_name, "w", zipfile.ZIP_DEFLATED) as zf:
        for fp in file_paths:
            if os.path.exists(fp):
                zf.write(fp, os.path.basename(fp))
    return zip_name


# ══════════════════════════════════════════════════════════════════════════════
# NAVIGATION
# ══════════════════════════════════════════════════════════════════════════════

is_admin = st.session_state.role == "admin"

with st.sidebar:
    st.markdown("""
    <div style="text-align:center;padding:18px 0 10px;">
        <div style="font-size:2.4rem;line-height:1;">🏫</div>
        <div style="color:#ffffff;font-weight:800;font-size:1.0rem;letter-spacing:.3px;margin-top:6px;">
            DHACSS Phase IV</div>
        <div style="color:rgba(255,255,255,.55);font-size:.73rem;letter-spacing:.5px;
                    text-transform:uppercase;margin-top:2px;">Timetable Pro</div>
    </div>""", unsafe_allow_html=True)
    st.divider()

    role_icon = "🛡️" if st.session_state.role == "admin" else "👁️"
    st.markdown(f"""
    <div style="text-align:center;margin-bottom:12px;">
        <span style="background:rgba(74,144,226,.25);border:1px solid rgba(74,144,226,.50);
                     border-radius:20px;padding:5px 16px;font-size:.82rem;
                     color:#d0e8ff;font-weight:600;letter-spacing:.3px;">
            {role_icon} &nbsp;{st.session_state.role.title()}
        </span>
    </div>""", unsafe_allow_html=True)

    st.markdown('<p class="section-label" style="color:rgba(255,255,255,.45);padding:0 4px;">Navigation</p>',
                unsafe_allow_html=True)

    pages = (["Dashboard","Configuration","Generate","Class View",
               "Teacher View","Analytics","Downloads"]
             if is_admin else ["Class View","Teacher View","Analytics"])
    menu = st.selectbox("Navigation", pages, label_visibility="collapsed")
    st.divider()
    if st.button("🚪 Logout", use_container_width=True):
        st.session_state.logged_in = False
        st.session_state.role = None
        st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════

if menu == "Dashboard":
    st.markdown("""
    <div style="background:linear-gradient(135deg,#f0f7ff 0%,#e8f0fb 100%);
                border:1px solid #bfdbfe;border-left:4px solid #4a90e2;
                border-radius:13px;padding:18px 24px;margin-bottom:22px;
                display:flex;align-items:center;gap:14px;">
        <div style="font-size:28px;">👋</div>
        <div>
            <div style="font-family:'Plus Jakarta Sans',sans-serif;font-weight:700;
                        font-size:16px;color:#1f4e79;margin-bottom:3px;">
                Welcome to School Scheduler Pro</div>
            <div style="color:#4a6fa5;font-size:13.5px;">
                Configure your school in <strong>Configuration</strong>, then click
                <strong>Generate</strong> to build the timetable automatically.
            </div>
        </div>
    </div>""", unsafe_allow_html=True)

    tt = st.session_state.timetable
    c1,c2,c3,c4 = st.columns(4)
    c1.metric("🏫 Sections",   len(st.session_state.sections))
    c2.metric("👩‍🏫 Teachers",  len(st.session_state.teachers))
    filled = sum(1 for sec in tt for day in DAYS for p in get_periods(day)
                 if p!="Lunch" and slot(sec,day,p)["subject"]) if tt else 0
    c3.metric("✅ Slots Filled", filled)
    clashes = len(validate_teacher_clashes()) if tt else 0
    c4.metric("⚠️ Clashes", clashes)

    st.markdown("""
    <div class="ui-card" style="margin-top:18px;border-left:4px solid #4a90e2;">
        <div style="display:flex;align-items:center;gap:12px;">
            <span style="font-size:1.8rem;">💡</span>
            <div>
                <div style="font-weight:700;color:#1f4e79;font-size:.95rem;">Getting Started</div>
                <div style="color:#5a7a9a;font-size:.88rem;margin-top:3px;">
                    Configure your school in <strong>Configuration</strong>,
                    then click <strong>Generate</strong> to build the timetable automatically.
                </div>
            </div>
        </div>
    </div>""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════════

if menu == "Configuration":
    t_sec, t_teach, t_ct, t_subj, t_assign = st.tabs(
        ["🏫 Sections","👩‍🏫 Teachers","🎓 Class Teachers","📚 Subjects","🔗 Assignments"])

    with t_sec:
        col1,col2 = st.columns(2)
        with col1:
            st.subheader("➕ Add Section")
            ns = st.text_input("Section name", key="ns")
            if st.button("Add Section", key="add_sec"):
                if ns.strip():
                    st.session_state.sections[ns.strip()]={};  save_all_data(); st.rerun()
                else: st.warning("Enter a section name.")
        with col2:
            st.subheader("🗑️ Remove Section")
            if st.session_state.sections:
                rs=st.selectbox("Select",list(st.session_state.sections),key="rs")
                if st.button("Delete Section", key="del_sec"):
                    for d in (st.session_state.sections, st.session_state.subject_config,
                              st.session_state.class_teachers): d.pop(rs,None)
                    save_all_data(); st.success(f"'{rs}' removed."); st.rerun()
        st.divider()
        if st.session_state.sections:
            st.markdown("**Current Sections:**")
            chips = "".join(
                f'<span style="background:#dbeafe;border:1px solid #bfdbfe;border-radius:8px;'
                f'padding:6px 14px;font-size:13px;font-weight:600;color:#1f4e79;'
                f'margin:3px;display:inline-block;">{sec}</span>'
                for sec in st.session_state.sections)
            st.markdown(f'<div style="margin-top:6px;">{chips}</div>', unsafe_allow_html=True)

    with t_teach:
        col1,col2 = st.columns(2)
        with col1:
            st.subheader("➕ Add Teacher")
            nt = st.text_input("Teacher name", key="nt")
            if st.button("Add Teacher", key="add_t"):
                if nt.strip():
                    st.session_state.teachers[clean(nt)]={};  save_all_data(); st.rerun()
                else: st.warning("Enter a teacher name.")
        with col2:
            st.subheader("🗑️ Remove Teacher")
            if st.session_state.teachers:
                rt=st.selectbox("Select",list(st.session_state.teachers),key="rt")
                if st.button("Delete Teacher", key="del_t"):
                    st.session_state.teachers.pop(rt,None); save_all_data()
                    st.success(f"'{rt}' removed."); st.rerun()
        st.divider()
        if st.session_state.teachers:
            st.markdown("**Current Teachers:**")
            chips = "".join(
                f'<span style="background:#f0fdf4;border:1px solid #bbf7d0;border-radius:8px;'
                f'padding:6px 14px;font-size:13px;font-weight:600;color:#065f46;'
                f'margin:3px;display:inline-block;">👩‍🏫 {t}</span>'
                for t in st.session_state.teachers)
            st.markdown(f'<div style="margin-top:6px;">{chips}</div>', unsafe_allow_html=True)

    with t_ct:
        if st.session_state.sections and st.session_state.teachers:
            ss=st.selectbox("Section",list(st.session_state.sections),key="ct_s")
            st2=st.selectbox("Teacher",list(st.session_state.teachers),key="ct_t")
            if st.button("Assign Class Teacher"):
                st.session_state.class_teachers[ss]=st2; save_all_data()
                st.success(f"{st2} assigned to {ss}."); st.rerun()
        else:
            st.info("Add sections and teachers first.")
        if st.session_state.class_teachers:
            st.divider()
            ct_df=pd.DataFrame(st.session_state.class_teachers.items(),columns=["Section","Class Teacher"])
            st.dataframe(ct_df,use_container_width=True,hide_index=True)
            rm_ct=st.selectbox("Remove CT for section",list(st.session_state.class_teachers),key="rm_ct")
            if st.button("Remove CT"):
                st.session_state.class_teachers.pop(rm_ct,None); save_all_data()
                st.success("Removed."); st.rerun()

    with t_subj:
        if not st.session_state.sections:
            st.info("Add sections first.")
        else:
            sel_sec=st.selectbox("Section",list(st.session_state.sections),key="ss_sec")
            col1,col2=st.columns(2)
            with col1:
                sn=st.text_input("Subject name",key="sn")
                wc=st.number_input("Weekly periods",min_value=1,max_value=10,step=1,key="wc")
                if st.button("Add Subject"):
                    if sn.strip():
                        st.session_state.subject_config.setdefault(sel_sec,{})[sn.strip()]=int(wc)
                        save_all_data(); st.success("Added."); st.rerun()
            with col2:
                cur_subjs=st.session_state.subject_config.get(sel_sec,{})
                if cur_subjs:
                    st.dataframe(pd.DataFrame(cur_subjs.items(),
                                              columns=["Subject","Periods/Week"]),
                                 hide_index=True,use_container_width=True)
                    rm_s=st.selectbox("Remove",list(cur_subjs),key="rm_s")
                    if st.button("Delete Subject"):
                        st.session_state.subject_config[sel_sec].pop(rm_s,None)
                        save_all_data(); st.success("Removed."); st.rerun()

    with t_assign:
        if not (st.session_state.teachers and st.session_state.subject_config):
            st.info("Add teachers and subjects first.")
        else:
            col1,col2=st.columns(2)
            with col1:
                st.subheader("➕ Assign")
                at=st.selectbox("Teacher",list(st.session_state.teachers),key="at")
                as_=st.selectbox("Section",list(st.session_state.subject_config),key="as_")
                ab_list=list(st.session_state.subject_config.get(as_,{}))
                if ab_list:
                    ab=st.selectbox("Subject",ab_list,key="ab")
                    if st.button("Assign"):
                        ta=st.session_state.teacher_assignment
                        ta.setdefault(at,{}).setdefault(as_,[])
                        if ab not in ta[at][as_]: ta[at][as_].append(ab)
                        save_all_data(); st.success("Assigned."); st.rerun()
            with col2:
                st.subheader("❌ Remove")
                ta=st.session_state.teacher_assignment
                if ta:
                    rm_at=st.selectbox("Teacher",list(ta),key="rm_at")
                    if rm_at in ta:
                        rm_as=st.selectbox("Section",list(ta[rm_at]),key="rm_as") if ta[rm_at] else None
                        if rm_as and rm_as in ta.get(rm_at,{}):
                            rm_ab=st.selectbox("Subject",ta[rm_at][rm_as],key="rm_ab")
                            if st.button("Remove Assignment"):
                                ta[rm_at][rm_as].remove(rm_ab)
                                if not ta[rm_at][rm_as]: del ta[rm_at][rm_as]
                                if not ta[rm_at]: del ta[rm_at]
                                save_all_data(); st.success("Removed."); st.rerun()
            st.divider()
            rows=[{"Teacher":t,"Section":s,"Subject":sub}
                  for t,secs in st.session_state.teacher_assignment.items()
                  for s,subs in secs.items() for sub in subs]
            if rows:
                st.dataframe(pd.DataFrame(rows),hide_index=True,use_container_width=True)

# ══════════════════════════════════════════════════════════════════════════════
# GENERATE
# ══════════════════════════════════════════════════════════════════════════════

if menu == "Generate":
    st.markdown("""
    <div style="background:linear-gradient(135deg,#f0f7ff 0%,#e0edff 100%);
                border:1px solid #bfdbfe;border-radius:13px;padding:20px 26px;margin-bottom:22px;">
        <div style="font-family:'Plus Jakarta Sans',sans-serif;font-size:18px;
                    font-weight:800;color:#1f4e79;margin-bottom:5px;">⚙️ Timetable Generator</div>
        <div style="color:#4a6fa5;font-size:13.5px;">
            Select the number of generation attempts — more runs improve quality by keeping
            the highest-fitness schedule across all attempts.
        </div>
    </div>""", unsafe_allow_html=True)

    RUNS = st.slider("Generation attempts (more = better quality)",
                     min_value=5, max_value=30, value=15, step=5)

    if st.button("🚀 Generate Timetable", key="gen_btn"):
        best_score, best_tt = -999_999, None
        progress = st.progress(0, text="Generating…")
        for run in range(RUNS):
            st.session_state.timetable = create_empty_timetable()
            assign_class_teacher_priority()
            assign_daily_singles()
            assign_math()
            assign_ix_x_doubles()
            basic_auto_fill()
            fill_under_quota_subjects()
            emergency_backfill()
            ensure_class_teacher_presence()
            force_fill()
            score = calculate_fitness()
            if score > best_score:
                best_score = score
                best_tt    = copy.deepcopy(st.session_state.timetable)
            progress.progress((run+1)/RUNS, text=f"Run {run+1}/{RUNS} — best score: {best_score}")
        st.session_state.timetable = best_tt
        save_all_data(); progress.empty()
        st.success(f"✅ Generation complete! Best fitness score: **{best_score}**")
        unmet = [(sec, subj, quota_remaining(sec, subj))
                 for sec, cfg in st.session_state.subject_config.items()
                 for subj in cfg if quota_remaining(sec, subj) > 0]
        if unmet:
            st.warning(f"⚠️ {len(unmet)} subject quota(s) could not be fully satisfied.")
            for sec, subj, rem in unmet[:10]:
                st.caption(f"  {sec} – {subj}: {rem} period(s) short")
        else:
            st.success("✅ All subject quotas fully met!")

    if st.session_state.timetable:
        st.divider()
        st.markdown("""
        <div style="background:#fffbeb;border:1px solid #fde68a;border-radius:12px;
                    padding:14px 20px;margin-bottom:14px;">
            <div style="font-family:'Plus Jakarta Sans',sans-serif;font-weight:700;
                        color:#92400e;font-size:14px;margin-bottom:2px;">🔄 Replace Teacher</div>
            <div style="color:#a16207;font-size:13px;">
                Swap all occurrences of one teacher with another across the entire timetable.
            </div>
        </div>""", unsafe_allow_html=True)

        c1,c2,c3 = st.columns([2,2,1])
        with c1: old_t=st.selectbox("Replace",list(st.session_state.teachers),key="rep_old")
        with c2: new_t=st.selectbox("With",   list(st.session_state.teachers),key="rep_new")
        with c3:
            st.write(""); st.write("")
            if st.button("Apply"):
                replace_teacher_everywhere(old_t,new_t)
                st.success(f"{old_t} → {new_t}")

        st.divider()
        st.markdown("""
        <div style="font-family:'Plus Jakarta Sans',sans-serif;font-size:16px;
                    font-weight:700;color:#1f4e79;margin-bottom:12px;">🔍 Validation Report</div>
        """, unsafe_allow_html=True)

        checks = {
            "3+ consecutive periods":   (validate_no_three_consecutive(), st.warning),
            "Unbalanced teacher loads": (validate_teacher_distribution(), st.warning),
            "Heavy Friday":             (validate_friday_load(),          st.info),
            "Maths non-consecutive":    (validate_maths_consecutive(),    st.info),
        }
        all_ok = True
        for label,(issues,fn) in checks.items():
            if issues:
                all_ok = False
                with st.expander(f"⚠️ {label} ({len(issues)})"):
                    for iss in issues: fn(iss)
        if all_ok: st.success("✅ No constraint violations found!")

# ══════════════════════════════════════════════════════════════════════════════
# CLASS VIEW
# ══════════════════════════════════════════════════════════════════════════════

if menu == "Class View":
    if not st.session_state.timetable:
        st.warning("Generate a timetable first.")
    else:
        sec = st.selectbox("Select Section", sorted(st.session_state.timetable.keys()), key="cv_sec")
        ct  = st.session_state.class_teachers.get(sec,"Not Assigned")

        st.markdown(f"""
        <div style="background:linear-gradient(135deg,#1f4e79 0%,#2b6cb0 100%);
                    border-radius:12px;padding:16px 24px;margin-bottom:18px;
                    display:flex;align-items:center;gap:16px;">
            <div style="background:rgba(255,255,255,0.15);border-radius:10px;
                        width:46px;height:46px;display:flex;align-items:center;
                        justify-content:center;font-size:22px;flex-shrink:0;">🏫</div>
            <div>
                <div style="color:rgba(255,255,255,0.65);font-size:10.5px;
                    font-family:'Plus Jakarta Sans',sans-serif;font-weight:600;
                    letter-spacing:1px;text-transform:uppercase;">Selected Section</div>
                <div style="color:#ffffff;font-size:20px;font-weight:800;
                    font-family:'Plus Jakarta Sans',sans-serif;">{sec}</div>
            </div>
            <div style="margin-left:auto;text-align:right;">
                <div style="color:rgba(255,255,255,0.65);font-size:10.5px;
                    font-family:'Plus Jakarta Sans',sans-serif;font-weight:600;
                    letter-spacing:1px;text-transform:uppercase;">Class Teacher</div>
                <div style="color:#ffffff;font-size:15px;font-weight:700;">{ct}</div>
            </div>
        </div>""", unsafe_allow_html=True)

        c1,c2,c3 = st.columns(3)
        c1.metric("Class Teacher", ct)
        c2.metric("Subjects", len(st.session_state.subject_config.get(sec,{})))
        c3.metric("Quota Gaps", sum(1 for s in st.session_state.subject_config.get(sec,{})
                                    if quota_remaining(sec,s)>0))

        df_data = {day:[
            (f"{slot(sec,day,p)['subject']}\n({slot(sec,day,p)['teacher']})"
             if slot(sec,day,p)["subject"] else "")
            if p in get_periods(day) else ""
            for p in ALL_PERIODS] for day in DAYS}

        df = pd.DataFrame({
            "Period":        ALL_PERIODS,
            "Mon-Wed Time":  [MON_WED_TIMES.get(p,"")  for p in ALL_PERIODS],
            "Monday":        df_data["Monday"],
            "Tuesday":       df_data["Tuesday"],
            "Wednesday":     df_data["Wednesday"],
            "Thu Time":      [THURSDAY_TIMES.get(p,"") for p in ALL_PERIODS],
            "Thursday":      df_data["Thursday"],
            "Friday":        df_data["Friday"],
            "Fri Time":      [FRIDAY_TIMES.get(p,"")   for p in ALL_PERIODS],
        }).set_index("Period")

        edited_df = st.data_editor(df, use_container_width=True,
                                   disabled=not is_admin, key="class_editor")

        col1,col2 = st.columns(2)
        with col1:
            if st.button("⬇ Download PDF", key="cv_pdf"):
                path = export_class_timetable_pdf(sec)
                with open(path,"rb") as f:
                    st.download_button("Download Class PDF", f, file_name=path,
                                       mime="application/pdf")
        with col2:
            if st.button("⬇ Download Word", key="cv_word"):
                path = export_class_timetable_to_word(sec)
                with open(path,"rb") as f:
                    st.download_button("Download Class Word", f, file_name=path,
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        if is_admin and st.button("💾 Save Manual Changes"):
            for day in DAYS:
                for period in ALL_PERIODS:
                    if period not in get_periods(day): continue
                    val = edited_df.loc[period, day] if day in edited_df.columns else ""
                    if val:
                        parts=str(val).split("\n"); subj=parts[0].strip()
                        teacher=(parts[1].replace("(","").replace(")","").strip().upper()
                                 if len(parts)>1 else "")
                    else:
                        subj=teacher=""
                    st.session_state.timetable[sec][day][period]={"subject":subj,"teacher":teacher}
            save_all_data(); st.success("Changes saved."); st.rerun()

        with st.expander("🔍 Validation", expanded=False):
            any_iss=False
            for iss in validate_subject_weekly(sec):     st.error(iss);   any_iss=True
            for iss in validate_teacher_clashes():       st.error(iss);   any_iss=True
            for iss in validate_class_teacher_presence():st.warning(iss); any_iss=True
            for iss in validate_teacher_max_load():      st.error(iss);   any_iss=True
            if not any_iss: st.success("No issues found.")

# ══════════════════════════════════════════════════════════════════════════════
# TEACHER VIEW
# ══════════════════════════════════════════════════════════════════════════════

if menu == "Teacher View":
    if not st.session_state.timetable:
        st.warning("Generate a timetable first.")
    else:
        teacher = st.selectbox("Select Teacher", sorted(st.session_state.teachers.keys()),
                               key="tv_teacher")
        total = count_teacher_periods(teacher)

        st.markdown(f"""
        <div style="background:linear-gradient(135deg,#1f4e79 0%,#2b6cb0 100%);
                    border-radius:12px;padding:16px 24px;margin-bottom:18px;
                    display:flex;align-items:center;gap:16px;">
            <div style="background:rgba(255,255,255,0.15);border-radius:10px;
                        width:46px;height:46px;display:flex;align-items:center;
                        justify-content:center;font-size:22px;flex-shrink:0;">👩‍🏫</div>
            <div>
                <div style="color:rgba(255,255,255,0.65);font-size:10.5px;
                    font-family:'Plus Jakarta Sans',sans-serif;font-weight:600;
                    letter-spacing:1px;text-transform:uppercase;">Teacher Schedule</div>
                <div style="color:#ffffff;font-size:20px;font-weight:800;
                    font-family:'Plus Jakarta Sans',sans-serif;">{teacher}</div>
            </div>
            <div style="margin-left:auto;text-align:right;">
                <div style="color:rgba(255,255,255,0.65);font-size:10.5px;
                    font-family:'Plus Jakarta Sans',sans-serif;font-weight:600;
                    letter-spacing:1px;text-transform:uppercase;">Weekly Load</div>
                <div style="color:#f59e0b;font-size:22px;font-weight:800;
                    font-family:'Plus Jakarta Sans',sans-serif;">{total} periods</div>
            </div>
        </div>""", unsafe_allow_html=True)

        c1,c2,c3 = st.columns(3)
        c1.metric("Total Weekly Periods", total)
        c2.metric("Avg / Day", f"{total/5:.1f}")
        busy = sum(1 for day in DAYS
                   if any(slot(sec,day,p)["teacher"]==teacher
                          for sec in st.session_state.timetable
                          for p in get_periods(day) if p!="Lunch"))
        c3.metric("Days Teaching", busy)

        df_data={}
        for day in DAYS:
            row=[]
            for p in ALL_PERIODS:
                if p=="Lunch":                row.append("Lunch"); continue
                if p not in get_periods(day): row.append(""); continue
                val=""
                for sec in st.session_state.timetable:
                    e=slot(sec,day,p)
                    if clean(e["teacher"])==clean(teacher): val=f"{sec}\n{e['subject']}"; break
                row.append(val)
            df_data[day]=row

        df=pd.DataFrame({
            "Period":        ALL_PERIODS,
            "Mon-Wed Time":  [MON_WED_TIMES.get(p,"")  for p in ALL_PERIODS],
            "Monday":        df_data["Monday"],
            "Tuesday":       df_data["Tuesday"],
            "Wednesday":     df_data["Wednesday"],
            "Thu Time":      [THURSDAY_TIMES.get(p,"") for p in ALL_PERIODS],
            "Thursday":      df_data["Thursday"],
            "Friday":        df_data["Friday"],
            "Fri Time":      [FRIDAY_TIMES.get(p,"")   for p in ALL_PERIODS],
        }).set_index("Period")
        st.dataframe(df, use_container_width=True)

        col1,col2=st.columns(2)
        with col1:
            path=export_teacher_view_word(teacher)
            with open(path,"rb") as f:
                st.download_button("⬇ Download Word", f, file_name=path,
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with col2:
            path=export_teacher_view_pdf(teacher)
            with open(path,"rb") as f:
                st.download_button("⬇ Download PDF", f, file_name=path,
                                   mime="application/pdf")

# ══════════════════════════════════════════════════════════════════════════════
# ANALYTICS
# ══════════════════════════════════════════════════════════════════════════════

if menu == "Analytics":
    if not st.session_state.timetable:
        st.warning("Generate a timetable first.")
    else:
        st.markdown("""<div style="font-family:'Plus Jakarta Sans',sans-serif;font-size:18px;
            font-weight:800;color:#1f4e79;margin-bottom:14px;">👩‍🏫 Teacher Workload Analysis</div>""",
            unsafe_allow_html=True)

        wl={t:count_teacher_periods(t) for t in st.session_state.teachers}
        wl_df=pd.DataFrame(wl.items(),columns=["Teacher","Total Periods"]) \
               .sort_values("Total Periods",ascending=False)

        def _wl(v):
            if v>=25: return "background-color:#ff6b6b"
            if v>=20: return "background-color:#ffd166"
            return "background-color:#90ee90"

        st.dataframe(wl_df.style.applymap(_wl,subset=["Total Periods"]),
                     use_container_width=True, hide_index=True)
        st.bar_chart(wl_df.set_index("Teacher"))

        st.markdown("""<div style="font-family:'Plus Jakarta Sans',sans-serif;font-size:18px;
            font-weight:800;color:#1f4e79;margin-top:24px;margin-bottom:14px;">
            📊 Quota Completeness</div>""", unsafe_allow_html=True)

        qrows=[]
        for sec,cfg in st.session_state.subject_config.items():
            for subj,target in cfg.items():
                placed=target-quota_remaining(sec,subj)
                qrows.append({"Section":sec,"Subject":subj,"Target":target,
                               "Placed":placed,"Remaining":max(0,target-placed)})
        if qrows:
            def _qc(v): return "background-color:#ff6b6b" if v>0 else "background-color:#90ee90"
            st.dataframe(pd.DataFrame(qrows).style.applymap(_qc,subset=["Remaining"]),
                         use_container_width=True, hide_index=True)

        st.markdown("""<div style="font-family:'Plus Jakarta Sans',sans-serif;font-size:18px;
            font-weight:800;color:#1f4e79;margin-top:24px;margin-bottom:14px;">
            📋 School Master Timetable</div>""", unsafe_allow_html=True)

        df_master=build_principal_matrix()
        st.dataframe(df_master, use_container_width=True)

        # =========================
        # ✅ USER INPUT (DEFINE FIRST)
        # =========================
        user_input = st.chat_input("Ask about timetable, swaps, teacher availability...")


        # =========================
        # ✅ FUNCTIONS
        # =========================

        def get_timetable_text():
            text = ""
            for sec in st.session_state.timetable:
                text += f"\nSection: {sec}\n"
                for day in DAYS:
                    text += f"{day}: "
                    for p in get_periods(day):
                        if p == "Lunch":
                            continue
                        cell = st.session_state.timetable[sec][day][p]
                        if cell["teacher"]:
                            text += f"{p}-{cell['teacher']} ({cell['subject']}), "
                    text += "\n"
            return text


        def get_all_teachers_view(df):
            text = ""

            # ✅ Normalize columns (fixes KeyError issue)
            df.columns = df.columns.str.strip().str.lower()

            # ✅ Required columns check
            required_cols = ['teacher', 'day', 'period', 'section']

            for col in required_cols:
                if col not in df.columns:
                    st.error(f"❌ Missing column: {col}")
                    st.write("Available columns:", df.columns.tolist())
                    return "Error in teacher data"

            teachers = df['teacher'].unique()

            for teacher in teachers:
                text += f"\n{teacher}:\n"

                teacher_df = df[df['teacher'] == teacher]
                grouped = teacher_df.groupby(['day', 'period'])['section'].apply(list)

                for (day, period), sections in grouped.items():
                    text += f"  {day} {period}: {', '.join(sections)}\n"

            return text


        # =========================
        # ✅ AI EXECUTION
        # =========================

        if user_input:
            with st.spinner("Thinking..."):

                timetable_data = get_timetable_text()
                teachers_data = get_all_teachers_view(df_master)

                st.write("📡 Sending THIS to AI:")

                full_prompt = f"""
        You are an expert school timetable optimizer and conflict resolver.

        Your task is to REMOVE conflicts for TABUSSUM (or any teacher mentioned in the user request).

        -----------------------
        USER REQUEST:
        {user_input}
        -----------------------

        FULL CLASS TIMETABLE:
        {timetable_data}

        -----------------------

        TEACHER TIMETABLE VIEW:
        {teachers_data}

        -----------------------

        STRICT RULES (MUST FOLLOW ALL):

        1. Focus ONLY on resolving the conflict mentioned.
        2. Do NOT create conflicts for any other teacher.
        3. A teacher can ONLY teach their assigned subject.
        4. Subject must NOT change (Physics stays Physics, etc).
        5. Do NOT change consecutive (double) periods.
        6. Swap ONLY with a teacher already assigned in the SAME time slot.
        7. Swap ONLY if subject matches.
        8. Make minimum possible change.
        9. Analyze BOTH timetable and teacher view before decision.

        -----------------------

        RESPONSE FORMAT:

        ✔ Conflict identified:
        ✔ Swap performed:
        ✔ Final result:

        If NOT possible:
        Say "Not possible" with reason.

        -----------------------

        IMPORTANT:
        - Be precise
        - Do NOT rewrite full timetable
        """

                # ✅ Debug view
                st.code(full_prompt)

                result = ask_ai(full_prompt)

            st.write("🤖 AI Response:")
            st.write(result)

        else:
            st.info("💡 Try: 'Remove Tabussum conflict on Monday Period 3'")


        if not df_master.empty:
            col1,col2=st.columns(2)
            with col1:
                path=export_timetable_word(df_master)
                with open(path,"rb") as f:
                    st.download_button("⬇ Download Word",f,file_name="school_timetable.docx")
            with col2:
                path=export_excel(df_master)
                with open(path,"rb") as f:
                    st.download_button("⬇ Download Excel",f,file_name="School_Timetable.xlsx")

# ══════════════════════════════════════════════════════════════════════════════
# DOWNLOADS
# ══════════════════════════════════════════════════════════════════════════════

if menu == "Downloads" and is_admin:
    if not st.session_state.timetable:
        st.warning("Generate a timetable first.")
    else:
        st.markdown("""
        <div style="background:linear-gradient(135deg,#f0f7ff 0%,#e0edff 100%);
                    border:1px solid #bfdbfe;border-radius:13px;
                    padding:20px 26px;margin-bottom:24px;">
            <div style="font-family:'Plus Jakarta Sans',sans-serif;font-size:18px;
                        font-weight:800;color:#1f4e79;margin-bottom:5px;">
                📦 Bulk Download Centre</div>
            <div style="color:#4a6fa5;font-size:13.5px;">
                Generate and download all timetable PDFs in a single ZIP file.
                Choose which set to include below.
            </div>
        </div>""", unsafe_allow_html=True)

        option = st.radio(
            "Select download set",
            ["All", "Class View Only", "Teacher View Only"],
            horizontal=True,
        )

        n_classes   = len(st.session_state.sections)
        n_teachers  = len(st.session_state.teachers)

        col1, col2, col3 = st.columns(3)
        col1.metric("📁 Class PDFs",   n_classes)
        col2.metric("📁 Teacher PDFs", n_teachers)
        col3.metric("📁 Total Files",
                    n_classes + n_teachers if option == "All"
                    else (n_classes if option == "Class View Only" else n_teachers))

        st.markdown("<br>", unsafe_allow_html=True)

        if st.button("⚙️ Prepare & Download ZIP", key="dl_zip_btn", use_container_width=True):
            class_files   = []
            teacher_files = []

            prog = st.progress(0, text="Generating PDFs…")
            total_items = n_classes + n_teachers
            done = 0

            if option in ("All", "Class View Only"):
                for section in st.session_state.sections:
                    path = export_class_timetable_pdf(section)
                    class_files.append(path)
                    done += 1
                    prog.progress(done / total_items,
                                  text=f"Class PDF: {section} ({done}/{total_items})")

            if option in ("All", "Teacher View Only"):
                for teacher in st.session_state.teachers:
                    path = export_teacher_view_pdf(teacher)
                    teacher_files.append(path)
                    done += 1
                    prog.progress(done / total_items,
                                  text=f"Teacher PDF: {teacher} ({done}/{total_items})")

            prog.empty()

            if option == "All":
                selected_files = class_files + teacher_files
                zip_name       = "all_timetables.zip"
            elif option == "Class View Only":
                selected_files = class_files
                zip_name       = "class_timetables.zip"
            else:
                selected_files = teacher_files
                zip_name       = "teacher_timetables.zip"

            if not selected_files:
                st.error("No files were generated. Ensure sections/teachers exist.")
            else:
                zip_path = create_zip(selected_files, zip_name)
                with open(zip_path, "rb") as zf:
                    st.download_button(
                        label=f"⬇ Download {zip_name}",
                        data=zf,
                        file_name=zip_name,
                        mime="application/zip",
                        use_container_width=True,
                    )
                st.success(f"✅ ZIP ready — {len(selected_files)} PDF(s) included.")

        st.divider()

        with st.expander("📄 Download individual class timetable PDF"):
            sec_dl = st.selectbox("Section", sorted(st.session_state.sections.keys()),
                                  key="dl_sec")
            if st.button("Generate & Download", key="dl_sec_btn"):
                path = export_class_timetable_pdf(sec_dl)
                with open(path,"rb") as f:
                    st.download_button("⬇ Download PDF", f, file_name=path,
                                       mime="application/pdf", key="dl_sec_dl")

        with st.expander("📄 Download individual teacher timetable PDF"):
            tch_dl = st.selectbox("Teacher", sorted(st.session_state.teachers.keys()),
                                  key="dl_tch")
            if st.button("Generate & Download", key="dl_tch_btn"):
                path = export_teacher_view_pdf(tch_dl)
                with open(path,"rb") as f:
                    st.download_button("⬇ Download PDF", f, file_name=path,
                                       mime="application/pdf", key="dl_tch_dl")
        st.markdown("### 📅 Day-wise Timetable")

        selected_day = st.selectbox("Select Day", DAYS, key="main_day_select")

        df_day = build_daywise_class_view(selected_day)

        st.dataframe(df_day)

        st.markdown("### 📚 Subject-wise Timetable")

        group = st.selectbox("Select Subject Group", list(SUBJECT_GROUPS.keys()))

        df_sub = build_subject_timetable(group)

        st.dataframe(df_sub)
########################### DAY WISE excel #############################


if st.button("Export Day-wise Excel"):
    path = export_daywise_excel(selected_day)
    st.success(f"File saved at {path}")
def export_daywise_excel(day):
    path = f"{day}_timetable.xlsx"

    df = build_daywise_class_view(day)

    # 🔥 Clean formatting (same as PDF)
    new_data = []

    for row in df.values:
        new_row = []
        for cell in row:
            if isinstance(cell, str) and "(" in cell:
                subject = cell.split("(")[0].strip()
                teacher = cell.split("(")[1].replace(")", "")
                short = teacher.split()[0][0] + "." + teacher.split()[-1]
                new_row.append(f"{subject}\n({short})")
            else:
                new_row.append(cell)
        new_data.append(new_row)

    df_clean = pd.DataFrame(new_data, columns=df.columns)

    # 🔥 Excel writer
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font

    wb = Workbook()
    ws = wb.active
    ws.title = day

    # 🔥 HEADINGS (same as you wanted)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(df.columns))
    ws.cell(row=1, column=1).value = "DHACSS PHASE IV CAMPUS"

    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(df.columns))
    ws.cell(row=2, column=1).value = "SENIOR SECTION TIMETABLE"

    ws.merge_cells(start_row=3, start_column=1, end_row=3, end_column=len(df.columns))
    ws.cell(row=3, column=1).value = "SESSION 2026–27"

    ws.merge_cells(start_row=4, start_column=1, end_row=4, end_column=len(df.columns))
    ws.cell(row=4, column=1).value = day.upper()

    # 🔥 Center align headings
    for r in range(1,5):
        cell = ws.cell(row=r, column=1)
        cell.alignment = Alignment(horizontal="center")
        cell.font = Font(bold=True)

    # 🔥 Table start row
    start_row = 6

    # headers
    for col, name in enumerate(df_clean.columns, 1):
        ws.cell(row=start_row, column=col).value = name

    # data
    for i, row in enumerate(df_clean.values, start_row+1):
        for j, val in enumerate(row, 1):
            cell = ws.cell(row=i, column=j)
            cell.value = val
            cell.alignment = Alignment(wrap_text=True, horizontal="center", vertical="center")

    # 🔥 Auto column width
    from openpyxl.utils import get_column_letter

    for col_num in range(1, ws.max_column + 1):
        col_letter = get_column_letter(col_num)
        max_len = 0

        for row in range(1, ws.max_row + 1):
            cell = ws.cell(row=row, column=col_num)
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))

        ws.column_dimensions[col_letter].width = min(max_len + 4, 30)

    wb.save(path)
    return path
if st.button("⬇ Download Day-wise Excel", use_container_width=True):
    path = export_daywise_excel(selected_day)
    with open(path, "rb") as f:
        st.download_button(
            label="📥 Click to Download Excel File",
            data=f,
            file_name=path,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
################## SUBJECT WISE PDF###################
def export_subjectwise_excel(group):
    path = f"{group}_timetable.xlsx"

    df = build_subject_timetable(group)

    # 🔥 Clean formatting (short teacher names + wrap)
    new_data = []
    for row in df.values:
        new_row = []
        for cell in row:
            if isinstance(cell, str) and "(" in cell:
                subject = cell.split("(")[0].strip()
                teacher = cell.split("(")[1].replace(")", "")
                short = teacher.split()[0][0] + "." + teacher.split()[-1]
                new_row.append(f"{subject}\n({short})")
            else:
                new_row.append(cell)
        new_data.append(new_row)

    df_clean = pd.DataFrame(new_data, columns=df.columns)

    # 🔥 Excel setup
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = group

    total_cols = len(df_clean.columns)

    # 🔥 HEADINGS
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=total_cols)
    ws.cell(row=1, column=1).value = "DHACSS PHASE IV CAMPUS"

    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=total_cols)
    ws.cell(row=2, column=1).value = "SENIOR SECTION"

    ws.merge_cells(start_row=3, start_column=1, end_row=3, end_column=total_cols)
    ws.cell(row=3, column=1).value = "SESSION 2026-27"

    ws.merge_cells(start_row=4, start_column=1, end_row=4, end_column=total_cols)
    ws.cell(row=4, column=1).value = f"SUBJECT ({group})"

    # 🔥 Heading style
    for r in range(1, 5):
        cell = ws.cell(row=r, column=1)
        cell.alignment = Alignment(horizontal="center")
        cell.font = Font(bold=True)

    # 🔥 Table start
    start_row = 6

    # headers
    for col, name in enumerate(df_clean.columns, 1):
        ws.cell(row=start_row, column=col).value = name

    # data
    for i, row in enumerate(df_clean.values, start_row + 1):
        for j, val in enumerate(row, 1):
            cell = ws.cell(row=i, column=j)
            cell.value = val
            cell.alignment = Alignment(
                wrap_text=True,
                horizontal="center",
                vertical="center"
            )

    # 🔥 Column width fix (safe version)
    for col_num in range(1, ws.max_column + 1):
        col_letter = get_column_letter(col_num)
        max_len = 0

        for row in range(1, ws.max_row + 1):
            cell = ws.cell(row=row, column=col_num)
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))

        ws.column_dimensions[col_letter].width = min(max_len + 4, 30)

    wb.save(path)
    return path

group = st.selectbox("Select Subject Group", list(SUBJECT_GROUPS.keys()))

df_sub = build_subject_timetable(group)

st.dataframe(df_sub)

# 👇 ab safe hai
path = export_subjectwise_excel(group)

with open(path, "rb") as f:
    st.download_button(
        label="⬇ Download Subject-wise Excel",
        data=f,
        file_name=path,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )